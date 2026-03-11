"""audit_equipment.py — Full arithmetic audit of equipment specs."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')

def parse_num(text):
    t = text.strip().replace(',','').replace('–','0').replace('-','0').replace(' ','')
    if not t or t == '0':
        return 0
    try:
        return int(t)
    except:
        return None

def audit_section(ti):
    """Audit a section table, return (items_sum, items_count, subtotal_from_doc)."""
    t = doc.tables[ti]
    section_sum = 0
    section_items = 0
    subtotal_val = None
    errors = []

    for ri in range(len(t.rows)):
        cells = [t.rows[ri].cells[ci].text.strip() for ci in range(min(len(t.columns), 7))]
        if len(cells) < 7:
            continue

        item_id = cells[0]
        # Item row
        if '.' in item_id and any(c.isdigit() for c in item_id) and 'รวม' not in cells[0]:
            qty = parse_num(cells[3])
            unit_price = parse_num(cells[5])
            total_val = parse_num(cells[6])

            if qty is not None and unit_price is not None and total_val is not None:
                expected = qty * unit_price
                ok = 'OK' if expected == total_val else f'MISMATCH (expect {expected:,})'
                if expected != total_val:
                    errors.append(f'T{ti} R{ri} {item_id}')
                name = cells[1][:42]
                print(f'    {item_id:6} {name:44} {qty:3} x {unit_price:>12,} = {total_val:>12,}  {ok}')
                section_sum += total_val
                section_items += 1

        # Subtotal row
        if 'รวม' in cells[0] or ('รวม' in cells[1]):
            sv = parse_num(cells[6])
            if sv is not None and sv > 0:
                subtotal_val = sv

    return section_sum, section_items, subtotal_val, errors


print('=' * 80)
print('FULL ARITHMETIC AUDIT')
print('=' * 80)

all_errors = []
grand_total = 0
grand_items = 0

# ═══ CATEGORY 1 ═══
print('\n' + '=' * 60)
print('CATEGORY 1: Network & Cybersecurity (15M target)')
print('=' * 60)
cat1_total = 0
cat1_items = 0

for ti, label in [(4, 'Network 1.1-1.7'), (5, 'Cybersecurity 2.1-2.7'), (6, 'Storage 3.1-3.3')]:
    print(f'\n  [{label}] Table {ti}')
    s, n, sub, errs = audit_section(ti)
    all_errors.extend(errs)
    if sub is not None:
        match = 'OK' if s == sub else f'MISMATCH (calc={s:,})'
        print(f'    --> Items sum: {s:,}  Subtotal in doc: {sub:,}  {match}')
        if s != sub:
            all_errors.append(f'T{ti} subtotal')
    cat1_total += s
    cat1_items += n

cat1_doc = parse_num(doc.tables[7].rows[0].cells[1].text)
match = 'OK' if cat1_total == cat1_doc else f'MISMATCH (calc={cat1_total:,})'
print(f'\n  CAT 1 TOTAL: {cat1_total:,}  Doc says: {cat1_doc:,}  {match}')
print(f'  CAT 1 ITEMS: {cat1_items}')
if cat1_total != cat1_doc:
    all_errors.append('Cat 1 total')
grand_total += cat1_total
grand_items += cat1_items


# ═══ CATEGORY 2 ═══
print('\n' + '=' * 60)
print('CATEGORY 2: AI Computing (67M target)')
print('=' * 60)
cat2_total = 0
cat2_items = 0

for ti, label in [(11, 'GPU Servers 4.1-4.4'), (13, 'CPU Servers 5.1-5.2'),
                  (14, 'Workstation & SW 6.1-6.5'), (15, 'DC Infra 7.1-7.5')]:
    print(f'\n  [{label}] Table {ti}')
    s, n, sub, errs = audit_section(ti)
    all_errors.extend(errs)
    if sub is not None:
        match = 'OK' if s == sub else f'MISMATCH (calc={s:,})'
        print(f'    --> Items sum: {s:,}  Subtotal in doc: {sub:,}  {match}')
        if s != sub:
            all_errors.append(f'T{ti} subtotal')
    cat2_total += s
    cat2_items += n

cat2_doc = parse_num(doc.tables[16].rows[0].cells[1].text)
# Note: T14 now has 7 rows (6.1-6.5 + header + subtotal) after DGX Spark insertion
match = 'OK' if cat2_total == cat2_doc else f'MISMATCH (calc={cat2_total:,})'
print(f'\n  CAT 2 TOTAL: {cat2_total:,}  Doc says: {cat2_doc:,}  {match}')
print(f'  CAT 2 ITEMS: {cat2_items}')
if cat2_total != cat2_doc:
    all_errors.append('Cat 2 total')
grand_total += cat2_total
grand_items += cat2_items


# ═══ CATEGORY 3 ═══
print('\n' + '=' * 60)
print('CATEGORY 3: Robotics Lab (68M target)')
print('=' * 60)
cat3_total = 0
cat3_items = 0

for ti, label in [(20, 'Cobots 8.1-8.6'), (21, 'Autonomous 9.1-9.5'),
                  (22, 'Sensing 10.1-10.6'), (23, 'Simulation 11.1-11.4'),
                  (24, 'Lab+Humanoid 12.1-12.13')]:
    print(f'\n  [{label}] Table {ti}')
    s, n, sub, errs = audit_section(ti)
    all_errors.extend(errs)
    if sub is not None:
        match = 'OK' if s == sub else f'MISMATCH (calc={s:,})'
        print(f'    --> Items sum: {s:,}  Subtotal in doc: {sub:,}  {match}')
        if s != sub:
            all_errors.append(f'T{ti} subtotal')
    cat3_total += s
    cat3_items += n

cat3_doc = parse_num(doc.tables[25].rows[0].cells[1].text)
match = 'OK' if cat3_total == cat3_doc else f'MISMATCH (calc={cat3_total:,})'
print(f'\n  CAT 3 TOTAL: {cat3_total:,}  Doc says: {cat3_doc:,}  {match}')
print(f'  CAT 3 ITEMS: {cat3_items}')
if cat3_total != cat3_doc:
    all_errors.append('Cat 3 total')
grand_total += cat3_total
grand_items += cat3_items


# ═══ GRAND TOTAL ═══
print('\n' + '=' * 60)
print('GRAND TOTAL')
print('=' * 60)

t28_total = parse_num(doc.tables[28].rows[16].cells[4].text)
match = 'OK' if grand_total == t28_total else f'MISMATCH (calc={grand_total:,})'
print(f'  Calculated:  {grand_total:,}')
print(f'  T28 R16:     {t28_total:,}  {match}')
print(f'  Total items: {grand_items}')
if grand_total != t28_total:
    all_errors.append('Grand total')

# T28 category cross-check
print('\n--- T28 Summary Table Cross-check ---')
t28 = doc.tables[28]
t28_c1 = parse_num(t28.rows[1].cells[4].text)
t28_c2 = parse_num(t28.rows[5].cells[4].text)
t28_c3 = parse_num(t28.rows[10].cells[4].text)
print(f'  Cat 1: T28={t28_c1:,}  calc={cat1_total:,}  {"OK" if t28_c1==cat1_total else "MISMATCH"}')
print(f'  Cat 2: T28={t28_c2:,}  calc={cat2_total:,}  {"OK" if t28_c2==cat2_total else "MISMATCH"}')
print(f'  Cat 3: T28={t28_c3:,}  calc={cat3_total:,}  {"OK" if t28_c3==cat3_total else "MISMATCH"}')
t28_sum = t28_c1 + t28_c2 + t28_c3
print(f'  Sum:   {t28_sum:,}  T28 total={t28_total:,}  {"OK" if t28_sum==t28_total else "MISMATCH"}')

# T28 section subtotals
print('\n--- T28 Section Subtotals ---')
section_checks = [
    (2, 'Network (1.1-1.7)', None),
    (3, 'Cyber (2.1-2.7)', None),
    (4, 'Storage (3.1-3.3)', None),
    (6, 'GPU (4.1-4.4)', None),
    (7, 'CPU (5.1-5.2)', None),
    (8, 'WS+SW (6.1-6.5)', None),
    (9, 'DC Infra (7.1-7.5)', None),
    (11, 'Cobots (8.1-8.6)', None),
    (12, 'Autonomous (9.1-9.5)', None),
    (13, 'Sensing (10.1-10.6)', None),
    (14, 'Simulation (11.1-11.4)', None),
    (15, 'Lab+Humanoid (12.1-12.13)', None),
]
for ri, label, _ in section_checks:
    val = parse_num(t28.rows[ri].cells[4].text)
    items = t28.rows[ri].cells[3].text.strip()
    print(f'  T28 R{ri:2}: {label:28} {items:>5}  {val:>12,}')

# T29 phasing audit
print('\n--- T29 Phasing Audit ---')
t29 = doc.tables[29]
for ri in range(1, 5):
    name = t29.rows[ri].cells[0].text.strip()[:25]
    c1 = parse_num(t29.rows[ri].cells[1].text)
    c2 = parse_num(t29.rows[ri].cells[2].text)
    c3 = parse_num(t29.rows[ri].cells[3].text)
    c4 = parse_num(t29.rows[ri].cells[4].text)
    row_sum = c1 + c2 + c3
    match = 'OK' if row_sum == c4 else f'MISMATCH (sum={row_sum:,})'
    print(f'  {name:27} {c1:>12,} + {c2:>12,} + {c3:>12,} = {row_sum:>12,}  doc={c4:>12,}  {match}')
    if row_sum != c4:
        all_errors.append(f'T29 R{ri} row sum')

# Column sums
for ci in range(1, 5):
    col_sum = sum(parse_num(t29.rows[ri].cells[ci].text) for ri in range(1, 4))
    total_row = parse_num(t29.rows[4].cells[ci].text)
    match = 'OK' if col_sum == total_row else f'MISMATCH (sum={col_sum:,})'
    print(f'  Col {ci} sum: {col_sum:,}  total row: {total_row:,}  {match}')
    if col_sum != total_row:
        all_errors.append(f'T29 Col {ci} sum')


# ═══ FINAL REPORT ═══
print('\n' + '=' * 60)
if all_errors:
    print(f'ERRORS FOUND: {len(all_errors)}')
    for e in all_errors:
        print(f'  - {e}')
else:
    print('ALL CHECKS PASSED - NO ERRORS')
print('=' * 60)
