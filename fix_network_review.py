"""
Network Expert Review & Fix — MFU-ARIC Equipment Specifications Budget
Fixes 14 issues across network, cybersecurity, and storage categories.
Budget-neutral: total remains exactly 15,000,000 THB.

Key changes:
  NETWORK (6,140K → 5,780K):
    1.1 Core Switch: 100G uplink spec update (same price)
    1.2 Access Switch: 1G→10G SFP+, remove PoE, qty 6→4 (1,500K→1,200K)
    1.3 IB Switch: HDR 200G→NDR 400G, model MQM8700→QM9700 (1,500K→1,800K)
    1.4 IB HCA: Spec update NDR 400G + PCIe 5.0 (same price)
    1.5 Fiber/DAC: OS2 single-mode→OM4 multi-mode + DAC cables (250K→300K)
    1.6 Wi-Fi: 6E→7, qty 12→8 (660K→520K)
    1.7 Replace Wireless Controller with OOB Management Switch (350K→80K)

  CYBERSECURITY (3,940K → 3,940K):
    2.1 NGFW: Clarify throughput spec (raw vs NGFW), add SSL VPN (same price)
    2.2 SIEM: Splunk→FortiAnalyzer+Wazuh, realistic pricing (1,200K→700K)
    2.3 EDR: CrowdStrike→FortiEDR, 200→100 endpoints (450K→350K)
    2.7 NEW: NAC moved from Phase 2 to Phase 1 (+600K)
    Fortinet Security Fabric consolidation for unified management

  STORAGE (4,920K → 5,280K, absorbs network savings):
    3.1 Primary: NetApp/Pure→HPE Alletra/Lenovo DE series (3,500K→3,800K)
    3.3 Cloud Backup: Increase scope (220K→280K)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

# ─── Helpers ──────────────────────────────────────────────────────

def set_cell_text(cell, text):
    """Replace all runs in the first paragraph, preserving first run's formatting."""
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')

def set_tr_cell(tr, col_idx, text):
    """Set text of a cell within a raw tr element."""
    tcs = tr.findall(qn('w:tc'))
    if col_idx >= len(tcs):
        return
    tc = tcs[col_idx]
    p = tc.find(qn('w:p'))
    if p is None:
        return
    runs = p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        p.remove(r)
    new_r = etree.SubElement(p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t_elem = etree.SubElement(new_r, qn('w:t'))
    t_elem.text = str(text)
    t_elem.set(qn('xml:space'), 'preserve')

def set_note_text(cell, text):
    """Replace text in a single-cell note table. Use | as paragraph separator."""
    parts = text.split('|')
    paras = cell.paragraphs
    for i, part in enumerate(parts):
        if i < len(paras):
            set_para_text(paras[i], part.strip())
        else:
            last_p = paras[-1]
            runs = last_p._p.findall(qn('w:r'))
            fmt = None
            if runs:
                rPr = runs[0].find(qn('w:rPr'))
                if rPr is not None:
                    fmt = deepcopy(rPr)
            new_r = etree.SubElement(last_p._p, qn('w:r'))
            if fmt is not None:
                new_r.insert(0, fmt)
            br = etree.SubElement(new_r, qn('w:br'))
            t_elem = etree.SubElement(new_r, qn('w:t'))
            t_elem.text = part.strip()
            t_elem.set(qn('xml:space'), 'preserve')
    for i in range(len(parts), len(paras)):
        cell._tc.remove(paras[i]._p)

def set_para_text(para, text):
    """Replace all runs in a paragraph, preserving first run's formatting."""
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t_elem = etree.SubElement(new_r, qn('w:t'))
    t_elem.text = str(text)
    t_elem.set(qn('xml:space'), 'preserve')

def fmt(n):
    """Format number with commas."""
    return f"{n:,}"


# ═══════════════════════════════════════════════════════════════════
print("Loading MFU_ARIC_Equipment_Specifications_Budget.docx ...")
doc = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')

# ─── TABLE 3: Rationale — update architecture description ─────────
print("  T3: Rationale (update architecture)")
set_note_text(doc.tables[3].rows[0].cells[0],
    "เหตุผลและความจำเป็น|"
    "ศูนย์ MFU-ARIC ต้องการระบบเครือข่ายที่มีความเร็วสูง (High-Performance Network) "
    "เพื่อรองรับการส่งข้อมูลขนาดใหญ่ระหว่าง GPU Cluster, Storage และ Workstation "
    "รวมถึงรองรับ AI Solutions Consulting Service และ Startup ที่ต้องเข้าถึง Computing Resources ได้อย่างปลอดภัย "
    "ระบบความปลอดภัยไซเบอร์ออกแบบตามแนวทาง Fortinet Security Fabric "
    "เพื่อการบริหารจัดการแบบรวมศูนย์และตอบสนองภัยคุกคามอัตโนมัติ "
    "ปกป้องทรัพย์สินทางปัญญา ข้อมูลภาคอุตสาหกรรม และข้อมูลธุรกิจของ Startup ที่เข้าร่วมโครงการ|"
    "สถาปัตยกรรมเครือข่ายแบ่งเป็น 3 ระดับ: "
    "(1) InfiniBand NDR 400Gbps Fabric สำหรับ GPU Cluster Interconnect "
    "(2) 10G/100G Ethernet Core Network สำหรับ Server-to-Storage และ Inter-VLAN Routing "
    "(3) Wi-Fi 7 Wireless Network สำหรับผู้ใช้งานทั่วไปและ IoT/Robotics|"
    "ระบบ Out-of-Band (OOB) Management Network แยกจากระบบ Production สำหรับบริหารจัดการเซิร์ฟเวอร์ผ่าน IPMI/BMC")


# ═══════════════════════════════════════════════════════════════════
# TABLE 4: Network Equipment (1.1–1.7)
# ═══════════════════════════════════════════════════════════════════
print("  T4: Network Equipment")
t4 = doc.tables[4]

# ─── 1.1 Core Switch: Update uplink to 100G ──────────────────────
print("    1.1 Core Switch: 40G→100G uplink")
set_cell_text(t4.rows[1].cells[2],
    "- 48 Port 10G SFP+ "
    "- 4 Port 100G QSFP28 Uplink "
    "- Layer 3 Managed, HA Redundant (VSS/StackWise Virtual) "
    "- Switching Capacity ≥ 1.2 Tbps "
    "- รองรับ SDN/VLAN/QoS/EVPN-VXLAN "
    "- Multi-tenant Isolation สำหรับ Startup แต่ละทีม")
# Price stays 600K × 2 = 1,200K

# ─── 1.2 Access Switch: 1G→10G, remove PoE, qty 6→4 ─────────────
print("    1.2 Access Switch: 1G→10G, qty 6→4")
set_cell_text(t4.rows[2].cells[1],
    "Access Switch 10G SFP+ (สำหรับ AI Server & Workstation)")
set_cell_text(t4.rows[2].cells[2],
    "- 24 Port 10G SFP+ "
    "- 4 Port 25G SFP28 Uplink "
    "- Layer 2/3, VLAN, QoS, ACL "
    "- Stackable, Non-blocking "
    "- Cisco Catalyst 9300-24S / Arista 7050SX3 หรือเทียบเท่า")
set_cell_text(t4.rows[2].cells[3], "4")
set_cell_text(t4.rows[2].cells[5], fmt(300000))
set_cell_text(t4.rows[2].cells[6], fmt(1200000))

# ─── 1.3 IB Switch: HDR→NDR 400G ─────────────────────────────────
print("    1.3 IB Switch: HDR 200G→NDR 400G")
set_cell_text(t4.rows[3].cells[1],
    "InfiniBand Switch NDR 400Gbps (สำหรับ GPU Cluster Interconnect)")
set_cell_text(t4.rows[3].cells[2],
    "- 32 Port NDR InfiniBand 400Gbps "
    "- Non-blocking Fabric "
    "- NVIDIA Quantum-2 QM9700 หรือเทียบเท่า "
    "- Latency < 90ns "
    "- รองรับ SHARP In-Network Computing")
set_cell_text(t4.rows[3].cells[5], fmt(1800000))
set_cell_text(t4.rows[3].cells[6], fmt(1800000))

# ─── 1.4 IB HCA: Update spec to NDR 400G + PCIe 5.0 ─────────────
print("    1.4 IB HCA: NDR 400G + PCIe 5.0")
set_cell_text(t4.rows[4].cells[1],
    "InfiniBand HCA Card NDR 400Gbps (ต่อพ่วงในแต่ละ GPU Server)")
set_cell_text(t4.rows[4].cells[2],
    "- NVIDIA Mellanox ConnectX-7 "
    "- 400Gbps NDR Single-port "
    "- PCIe 5.0 x16 "
    "- RDMA/GPUDirect Support")
# Price stays 85K × 8 = 680K

# ─── 1.5 Fiber: OS2→OM4 + DAC cables ─────────────────────────────
print("    1.5 Fiber: OS2→OM4 + DAC cables")
set_cell_text(t4.rows[5].cells[1],
    "Fiber Optic Cable, DAC Cable & Patch Panel")
set_cell_text(t4.rows[5].cells[2],
    "- OM4 Multi-mode Fiber (50/125μm) สำหรับ 10G/25G/100G Ethernet "
    "- LC-LC Duplex 5m/10m/20m "
    "- DAC (Direct Attach Copper) 400G 1m/2m/3m สำหรับ InfiniBand "
    "- 24-Port Patch Panel ×4 "
    "- Cable Tray & Management")
set_cell_text(t4.rows[5].cells[5], fmt(300000))
set_cell_text(t4.rows[5].cells[6], fmt(300000))

# ─── 1.6 Wi-Fi: 6E→7, qty 12→8 ──────────────────────────────────
print("    1.6 Wi-Fi AP: 6E→7, qty 12→8")
set_cell_text(t4.rows[6].cells[1],
    "Wi-Fi 7 Access Point (Cisco Catalyst 9178 หรือเทียบเท่า)")
set_cell_text(t4.rows[6].cells[2],
    "- 802.11be Tri-band (2.4/5/6 GHz) "
    "- Speed ≥ 11 Gbps aggregate "
    "- WPA3, MU-MIMO 16×16, MLO (Multi-Link Operation) "
    "- Cloud Managed (รวม License 3 ปี)")
set_cell_text(t4.rows[6].cells[3], "8")
set_cell_text(t4.rows[6].cells[5], fmt(65000))
set_cell_text(t4.rows[6].cells[6], fmt(520000))

# ─── 1.7 Replace Wireless Controller → OOB Management Switch ─────
print("    1.7 Wireless Controller→OOB Management Switch")
set_cell_text(t4.rows[7].cells[0], "1.7")
set_cell_text(t4.rows[7].cells[1],
    "Out-of-Band Management Switch (OOB)")
set_cell_text(t4.rows[7].cells[2],
    "- 24 Port 1G RJ45 Managed "
    "- Layer 2, VLAN "
    "- สำหรับ IPMI/iDRAC/BMC Management Network "
    "- แยก Network จาก Production Traffic")
set_cell_text(t4.rows[7].cells[3], "1")
set_cell_text(t4.rows[7].cells[4], "เครื่อง")
set_cell_text(t4.rows[7].cells[5], fmt(80000))
set_cell_text(t4.rows[7].cells[6], fmt(80000))

# ─── Network subtotal ────────────────────────────────────────────
set_cell_text(t4.rows[8].cells[6], fmt(5780000))
print(f"    Network subtotal: 5,780,000")


# ═══════════════════════════════════════════════════════════════════
# TABLE 5: Cybersecurity (2.1–2.7)
# ═══════════════════════════════════════════════════════════════════
print("  T5: Cybersecurity")
t5 = doc.tables[5]

# ─── 2.1 NGFW: Clarify throughput spec, add SSL VPN ──────────────
print("    2.1 NGFW: Clarify throughput (raw vs NGFW)")
set_cell_text(t5.rows[1].cells[2],
    "- Firewall Throughput ≥ 65 Gbps (ASIC-accelerated) "
    "- NGFW Throughput (IPS + App Control) ≥ 10 Gbps "
    "- SSL Inspection Throughput ≥ 7 Gbps "
    "- IPS, Application Control, Web Filtering "
    "- HA Active-Passive Redundant "
    "- SSL VPN (≥ 500 Concurrent Users สำหรับ Remote Access) "
    "- Threat Intelligence Auto-Update "
    "- รวม License 3 ปี (FortiGuard UTP Bundle)")
# Price stays 850K × 2 = 1,700K

# ─── 2.2 SIEM: Splunk→FortiAnalyzer+Wazuh ────────────────────────
print("    2.2 SIEM: Splunk→FortiAnalyzer+Wazuh (realistic pricing)")
set_cell_text(t5.rows[2].cells[1],
    "SIEM & Log Analytics Platform (FortiAnalyzer + Wazuh)")
set_cell_text(t5.rows[2].cells[2],
    "- FortiAnalyzer VM/HW สำหรับ Fortinet Device Log Analytics "
    "- Wazuh SIEM/XDR (Open-source, self-hosted) "
    "- Log Ingestion ≥ 100 GB/day "
    "- Real-time Threat Detection & Compliance Reporting "
    "- Integration กับ Fortinet Security Fabric "
    "- License 3 ปี (FortiAnalyzer)")
set_cell_text(t5.rows[2].cells[5], fmt(700000))
set_cell_text(t5.rows[2].cells[6], fmt(700000))

# ─── 2.3 EDR: CrowdStrike→FortiEDR, 200→100 endpoints ────────────
print("    2.3 EDR: CrowdStrike→FortiEDR, 200→100 endpoints")
set_cell_text(t5.rows[3].cells[1],
    "Endpoint Detection & Response (EDR) (FortiEDR หรือเทียบเท่า)")
set_cell_text(t5.rows[3].cells[2],
    "- Agent-based สำหรับ Endpoint 100 เครื่อง "
    "- AI-driven Threat Hunting & Zero-Day Protection "
    "- Integration กับ FortiGate & FortiAnalyzer (Security Fabric) "
    "- License 3 ปี")
set_cell_text(t5.rows[3].cells[5], fmt(350000))
set_cell_text(t5.rows[3].cells[6], fmt(350000))

# R4 (SOC WS), R5 (UPS), R6 (Vuln Scanner) — keep as-is

# ─── 2.7 NEW: NAC (moved from Phase 2 to Phase 1) ────────────────
print("    2.7 NAC: Insert new row (moved from Phase 2)")
# Clone row 6 (item 2.6) and insert after it
source_tr = t5.rows[6]._tr
new_tr = deepcopy(source_tr)
source_tr.addnext(new_tr)

# Set cells of the new row (2.7 NAC)
set_tr_cell(new_tr, 0, "2.7")
set_tr_cell(new_tr, 1,
    "Network Access Control (NAC) (FortiNAC หรือเทียบเท่า)")
set_tr_cell(new_tr, 2,
    "- Endpoint Profiling & Classification "
    "- 802.1X Authentication "
    "- Guest/Startup Portal (Multi-tenant Access Control) "
    "- Integration กับ FortiGate & FortiAnalyzer "
    "- License 100 Endpoints, 3 ปี")
set_tr_cell(new_tr, 3, "1")
set_tr_cell(new_tr, 4, "ชุด")
set_tr_cell(new_tr, 5, fmt(600000))
set_tr_cell(new_tr, 6, fmt(600000))

# ─── Update subtotal row (now pushed to last position) ────────────
# After inserting new row, the subtotal is the last row
subtotal_tr = t5._tbl.findall(qn('w:tr'))[-1]
# Update the reference text and total
set_tr_cell(subtotal_tr, 0, "รวมระบบความปลอดภัยไซเบอร์ ระยะที่ 1 (2.1–2.7)")
# Set all merged cells in the subtotal row to the same text
for ci in range(1, 6):
    set_tr_cell(subtotal_tr, ci, "รวมระบบความปลอดภัยไซเบอร์ ระยะที่ 1 (2.1–2.7)")
set_tr_cell(subtotal_tr, 6, fmt(3940000))
print(f"    Cyber subtotal: 3,940,000 (unchanged)")


# ═══════════════════════════════════════════════════════════════════
# TABLE 6: Storage (3.1–3.3)
# ═══════════════════════════════════════════════════════════════════
print("  T6: Storage")
t6 = doc.tables[6]

# ─── 3.1 Primary: NetApp/Pure→HPE/Lenovo, price 3.5M→3.8M ───────
print("    3.1 Primary Storage: NetApp/Pure→HPE Alletra/Lenovo DE")
set_cell_text(t6.rows[1].cells[1],
    "All-Flash SSD/NVMe Storage (Primary) (HPE Alletra 5000 / Lenovo ThinkSystem DE6400 หรือเทียบเท่า)")
set_cell_text(t6.rows[1].cells[2],
    "- ความจุ Usable ≥ 200 TB\n"
    "- IOPS ≥ 500,000 (Mixed R/W)\n"
    "- NVMe/SAS SSD with NVMe Cache Tier\n"
    "- Deduplication & Compression (Inline)\n"
    "- HA Dual Controller, 25GbE/100GbE Interface\n"
    "- Snapshot & Thin Provisioning")
set_cell_text(t6.rows[1].cells[5], fmt(3800000))
set_cell_text(t6.rows[1].cells[6], fmt(3800000))

# ─── 3.3 Cloud Backup: Increase budget 220K→280K ─────────────────
print("    3.3 Cloud Backup: 220K→280K")
set_cell_text(t6.rows[3].cells[2],
    "- Cloud-based Disaster Recovery\n"
    "- Automated Backup Schedule (Daily Incremental, Weekly Full)\n"
    "- Geo-redundant Storage (≥ 10 TB Critical Data)\n"
    "- Integration กับ On-premise Storage\n"
    "- License 3 ปี")
set_cell_text(t6.rows[3].cells[5], fmt(280000))
set_cell_text(t6.rows[3].cells[6], fmt(280000))

# ─── Storage subtotal ────────────────────────────────────────────
set_cell_text(t6.rows[4].cells[6], fmt(5280000))
print(f"    Storage subtotal: 5,280,000")


# ═══════════════════════════════════════════════════════════════════
# TABLE 7: Category 1 Grand Total — stays 15,000,000
# ═══════════════════════════════════════════════════════════════════
# No change needed — total is still 15M
print(f"  T7: Grand total: 15,000,000 (unchanged)")


# ═══════════════════════════════════════════════════════════════════
# TABLE 8: Notes — complete rewrite
# ═══════════════════════════════════════════════════════════════════
print("  T8: Notes (rewrite)")
set_note_text(doc.tables[8].rows[0].cells[0],
    "หมายเหตุ: "
    "| ราคา License ที่ระบุเป็นราคา 3 ปีแรก (FortiGuard UTP, FortiAnalyzer, FortiEDR, FortiNAC, Cloud Backup) "
    "หลังจากนั้นจะจ่ายเป็นค่าต่ออายุในงบดำเนินการ ประมาณ 1.0 ล้านบาท/ปี"
    "| InfiniBand Switch (NDR 400Gbps) และ HCA (ConnectX-7 NDR) เป็นอุปกรณ์จำเป็นสำหรับ GPU Cluster High-Bandwidth Interconnect "
    "ที่มีความหน่วงต่ำ จำนวน HCA 8 การ์ดรวมสำรอง 1 การ์ดสำหรับ Hot Standby"
    "| ออกแบบระบบ Cybersecurity ตามแนวทาง Fortinet Security Fabric (FortiGate + FortiAnalyzer + FortiEDR + FortiNAC) "
    "เพื่อการบริหารจัดการแบบรวมศูนย์ ลดความซับซ้อนในการปฏิบัติงาน และเพิ่มประสิทธิภาพการตอบสนองภัยคุกคามอัตโนมัติ"
    "| ย้ายรายการ NAC จากระยะที่ 2 มาเป็นระยะที่ 1 เนื่องจากจำเป็นต้องใช้ตั้งแต่เริ่มเปิดศูนย์ฯ "
    "สำหรับการควบคุมการเข้าถึงเครือข่ายแบบ Multi-tenant (แยก Startup แต่ละทีม)"
    "| รายการ Cybersecurity ระยะที่ 2 (PAM, DLP, Cyber Range) วงเงินประมาณ 1,550,000 บาท "
    "จะจัดซื้อเพิ่มเติมในปีที่ 3–4 จากงบดำเนินการหรืองบสำรอง"
    "| เปลี่ยน Wireless Controller เป็น Cloud-managed License (รวมอยู่ในราคา AP แล้ว) "
    "และเพิ่ม OOB Management Switch สำหรับ IPMI/BMC Management แยกจาก Production Network"
    "| ใช้ Cloud Backup (รายการ 3.3) สำหรับ Critical Data ≥ 10 TB (Config, Database, Source Code, IP) "
    "จัดเป็นงบลงทุนเนื่องจากเป็นการจ่ายล่วงหน้า 3 ปี (Prepaid Subscription)"
    "| ค่าบำรุงรักษาอุปกรณ์เครือข่ายและ Security ปีที่ 2–5 ประมาณ 1.2–1.5 ล้านบาท/ปี จ่ายจากงบดำเนินการ")


# ═══════════════════════════════════════════════════════════════════
# TABLE 28: Summary table — update item counts and descriptions
# ═══════════════════════════════════════════════════════════════════
print("  T28: Summary table")
t28 = doc.tables[28]

# R1: Total items 16→17
set_cell_text(t28.rows[1].cells[3], "17 รายการ")

# R2: Network description
set_cell_text(t28.rows[2].cells[2], "Switch 10G/100G, InfiniBand NDR 400G, Wi-Fi 7, OOB Mgmt")
set_cell_text(t28.rows[2].cells[4], fmt(5780000))

# R3: Cyber description and item count
set_cell_text(t28.rows[3].cells[2], "Firewall, SIEM, EDR, SOC, NAC (Fortinet Fabric)")
set_cell_text(t28.rows[3].cells[3], "7")
set_cell_text(t28.rows[3].cells[4], fmt(3940000))

# R4: Storage
set_cell_text(t28.rows[4].cells[2], "NVMe/SSD Storage, NAS, Cloud Backup")
set_cell_text(t28.rows[4].cells[4], fmt(5280000))


# ═══════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════
outfile = 'MFU_ARIC_Equipment_Specifications_Budget.docx'
doc.save(outfile)
print(f"\nSaved: {outfile}")

print("\n" + "="*60)
print("SUMMARY OF CHANGES")
print("="*60)
print("""
NETWORK (6,140K → 5,780K, saves 360K):
  1.1 Core Switch      1,200K → 1,200K  (100G uplink spec update)
  1.2 Access Switch     1,500K → 1,200K  (1G→10G SFP+, qty 6→4)
  1.3 IB Switch         1,500K → 1,800K  (HDR 200G→NDR 400G)
  1.4 IB HCA              680K →   680K  (NDR 400G + PCIe 5.0 spec)
  1.5 Fiber/DAC            250K →   300K  (OS2→OM4 + DAC cables)
  1.6 Wi-Fi AP             660K →   520K  (6E→7, qty 12→8)
  1.7 OOB Mgmt Switch      350K →    80K  (replaced wireless controller)

CYBERSECURITY (3,940K → 3,940K, budget-neutral):
  2.1 NGFW              1,700K → 1,700K  (clarify throughput, add SSL VPN)
  2.2 SIEM              1,200K →   700K  (Splunk→FortiAnalyzer+Wazuh)
  2.3 EDR                 450K →   350K  (CrowdStrike→FortiEDR, 100 EP)
  2.4 SOC WS              240K →   240K  (no change)
  2.5 UPS                 170K →   170K  (no change)
  2.6 Vuln Scanner        180K →   180K  (no change)
  2.7 NAC                 NEW  →   600K  (moved from Phase 2)

STORAGE (4,920K → 5,280K, absorbs 360K savings):
  3.1 Primary Storage   3,500K → 3,800K  (HPE/Lenovo, realistic pricing)
  3.2 NAS               1,200K → 1,200K  (no change)
  3.3 Cloud Backup        220K →   280K  (scope clarified, ≥10TB critical)

GRAND TOTAL: 15,000,000 (unchanged)
""")
