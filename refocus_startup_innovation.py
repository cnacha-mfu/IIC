"""
Refocus both MFU-ARIC proposal documents toward startup/business
and innovation development support rather than research/publication.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

# ─── Helper ───────────────────────────────────────────────────────
def set_cell_text(cell, text):
    """Replace all runs in the first paragraph of a cell, preserving formatting from the first run."""
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    # Copy formatting from the first existing run if available
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    # Remove all existing runs
    for r in runs:
        para._p.remove(r)
    # Create new run with preserved formatting
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')

def set_para_text(para, text):
    """Replace all runs in a paragraph, preserving formatting from the first run."""
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')

def set_note_text(cell, text):
    """Replace text in a single-cell note table. Use | as paragraph separator."""
    parts = text.split('|')
    paras = cell.paragraphs
    # Reuse existing paragraphs or add to last one
    for i, part in enumerate(parts):
        if i < len(paras):
            set_para_text(paras[i], part.strip())
        else:
            # Just append to last paragraph with newline
            last_p = paras[-1]
            runs = last_p._p.findall(qn('w:r'))
            fmt = None
            if runs:
                rPr = runs[0].find(qn('w:rPr'))
                if rPr is not None:
                    fmt = deepcopy(rPr)
            new_r = etree.SubElement(last_p._p, qn('w:r'))
            if fmt is not None:
                new_r.insert(0, fmt)
            br = etree.SubElement(new_r, qn('w:br'))
            t = etree.SubElement(new_r, qn('w:t'))
            t.text = part.strip()
            t.set(qn('xml:space'), 'preserve')
    # Remove excess paragraphs
    for i in range(len(parts), len(paras)):
        cell._tc.remove(paras[i]._p)


# ═══════════════════════════════════════════════════════════════════
# PART 1: MFU_ARIC_Government_Funding_Proposal_2569.docx
# ═══════════════════════════════════════════════════════════════════
print("Loading Government Funding Proposal...")
doc = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# ─── TABLE 0: Project summary ────────────────────────────────────
print("  Table 0: Project type")
set_cell_text(doc.tables[0].rows[0].cells[1],
    "โครงการลงทุน (Capital Investment) ด้านโครงสร้างพื้นฐาน นวัตกรรม และพัฒนากำลังคนดิจิทัล")

# ─── TABLE 1: Basic info ─────────────────────────────────────────
print("  Table 1: Project type")
set_cell_text(doc.tables[1].rows[3].cells[1],
    "โครงการพัฒนา (Development Project) ด้านโครงสร้างพื้นฐาน นวัตกรรม และพัฒนากำลังคนดิจิทัล")

# ─── TABLE 2: Executive summary box ──────────────────────────────
print("  Table 2: Executive summary")
set_note_text(doc.tables[2].rows[0].cells[0],
    "ความเร่งด่วนและความจำเป็น: ประเทศไทยกำลังเผชิญกับการขาดแคลนบุคลากรด้าน AI และ Robotics อย่างเฉียบพลัน ขณะที่ประเทศคู่แข่งในอาเซียน เช่น สิงคโปร์ มาเลเซีย และเวียดนาม ต่างลงทุนอย่างมหาศาลในการสร้างศูนย์นวัตกรรมและบ่มเพาะ Startup ด้าน AI โครงการนี้จะช่วยให้ประเทศไทยสามารถสร้างระบบนิเวศ AI Startup สร้างธุรกิจนวัตกรรม และลดการพึ่งพาเทคโนโลยีต่างชาติในอุตสาหกรรมอนาคต|"
    "ความเหมาะสมของมหาวิทยาลัยแม่ฟ้าหลวง: มฟล. มีความพร้อมด้านวิชาการ เครือข่ายนวัตกรรม ส่วนจัดการทรัพย์สินทางปัญญาและนวัตกรรม (MFII) ที่มีประสบการณ์บ่มเพาะ Startup และที่ตั้งเชิงยุทธศาสตร์ที่ประตูสู่ GMS อีกทั้งมีสำนักวิชาเทคโนโลยีดิจิทัลประยุกต์ที่มีกิจกรรม AI และนวัตกรรมต่อเนื่อง")

# ─── TABLE 4: KPIs — major restructure ───────────────────────────
print("  Table 4: KPIs (restructure)")
t4 = doc.tables[4]

# R1: Section header — rename from research to applied innovation
set_cell_text(t4.rows[1].cells[0], "OUTPUT – ด้านนวัตกรรมและวิจัยประยุกต์")
set_cell_text(t4.rows[1].cells[1], "OUTPUT – ด้านนวัตกรรมและวิจัยประยุกต์")
set_cell_text(t4.rows[1].cells[2], "OUTPUT – ด้านนวัตกรรมและวิจัยประยุกต์")
set_cell_text(t4.rows[1].cells[3], "OUTPUT – ด้านนวัตกรรมและวิจัยประยุกต์")

# R2: Reduce publications — focus on applied/industry papers
set_cell_text(t4.rows[2].cells[0], "บทความวิจัยประยุกต์/Industry Paper (สะสม)")
set_cell_text(t4.rows[2].cells[1], "5 บทความ")
set_cell_text(t4.rows[2].cells[2], "15 บทความ")
set_cell_text(t4.rows[2].cells[3], "30 บทความ")

# R3: Patents — increase significantly (core IP output)
set_cell_text(t4.rows[3].cells[0], "สิทธิบัตร/อนุสิทธิบัตร/Software License (สะสม)")
set_cell_text(t4.rows[3].cells[1], "5 รายการ")
set_cell_text(t4.rows[3].cells[2], "15 รายการ")
set_cell_text(t4.rows[3].cells[3], "30 รายการ")

# R4: Industry collaboration projects — increase
set_cell_text(t4.rows[4].cells[0], "โครงการร่วมพัฒนากับภาคเอกชน/SME (สะสม)")
set_cell_text(t4.rows[4].cells[1], "8 โครงการ")
set_cell_text(t4.rows[4].cells[2], "25 โครงการ")
set_cell_text(t4.rows[4].cells[3], "50 โครงการ")

# R9: Section header — elevate innovation/business
set_cell_text(t4.rows[9].cells[0], "OUTPUT – ด้าน Startup/ธุรกิจนวัตกรรม")
set_cell_text(t4.rows[9].cells[1], "OUTPUT – ด้าน Startup/ธุรกิจนวัตกรรม")
set_cell_text(t4.rows[9].cells[2], "OUTPUT – ด้าน Startup/ธุรกิจนวัตกรรม")
set_cell_text(t4.rows[9].cells[3], "OUTPUT – ด้าน Startup/ธุรกิจนวัตกรรม")

# R10: Startups incubated — major increase
set_cell_text(t4.rows[10].cells[0], "Startup/ธุรกิจนวัตกรรมที่บ่มเพาะ (สะสม)")
set_cell_text(t4.rows[10].cells[1], "8")
set_cell_text(t4.rows[10].cells[2], "25")
set_cell_text(t4.rows[10].cells[3], "50")

# R11: Revenue — increase targets
set_cell_text(t4.rows[11].cells[0], "รายได้จากบริการ/ธุรกิจนวัตกรรม (ล้านบาท/ปี)")
set_cell_text(t4.rows[11].cells[1], "10")
set_cell_text(t4.rows[11].cells[2], "30")
set_cell_text(t4.rows[11].cells[3], "60")

# R14: Economic value — increase
set_cell_text(t4.rows[14].cells[0], "มูลค่าทางเศรษฐกิจที่สร้างจาก Startup/นวัตกรรม (ล้านบาท สะสม)")
set_cell_text(t4.rows[14].cells[1], "100")
set_cell_text(t4.rows[14].cells[2], "500")
set_cell_text(t4.rows[14].cells[3], "2,000")

# ─── TABLE 5: 4 Hubs — reorder and reframe ───────────────────────
print("  Table 5: Reorder Hubs (Innovation first)")
t5 = doc.tables[5]

# Swap R1 (Research) and R3 (Innovation) — make Innovation first
# Also rewrite Research Hub to be "Applied Research"
# Row 1: Innovation Hub (was row 3)
set_cell_text(t5.rows[1].cells[0], "Innovation Hub (ศูนย์นวัตกรรมและธุรกิจ)")
set_cell_text(t5.rows[1].cells[1], "Startup Incubation & Acceleration, Technology Transfer, Industry Co-development, AI Solutions Consulting, Hackathon, Demo Day, IP Commercialization")
set_cell_text(t5.rows[1].cells[2], "Startup, ธุรกิจนวัตกรรม, IP License, รายได้เชิงพาณิชย์, AI Solutions")
set_cell_text(t5.rows[1].cells[3], "ผู้ประกอบการ, SME, นักศึกษา, ภาคอุตสาหกรรม")

# Row 2: Education Hub (stays)
# (no change needed)

# Row 3: Applied Research Hub (was row 1, reframed)
set_cell_text(t5.rows[3].cells[0], "Applied Research Hub (ศูนย์วิจัยประยุกต์)")
set_cell_text(t5.rows[3].cells[1], "วิจัยประยุกต์ร่วมภาคอุตสาหกรรม (AI for Health, Agri-tech, Smart City, GMS Languages); พัฒนาต้นแบบผลิตภัณฑ์; ถ่ายทอดเทคโนโลยีสู่ SME")
set_cell_text(t5.rows[3].cells[2], "ต้นแบบผลิตภัณฑ์, สิทธิบัตร, AI Solutions สำหรับอุตสาหกรรม")
set_cell_text(t5.rows[3].cells[3], "ภาคอุตสาหกรรม, SME, บัณฑิตศึกษา")

# Row 4: Regional Hub — add business angle
set_cell_text(t5.rows[4].cells[1], "MOU กับ CLMV, GMS Business Network, Cross-border AI Trade, AI Diplomacy, แลกเปลี่ยนผู้ประกอบการ")

# ─── TABLE 6: Work plan — shift toward startup/innovation ────────
print("  Table 6: Work plan (startup-focused)")
t6 = doc.tables[6]

set_cell_text(t6.rows[1].cells[2],
    "• จัดตั้งโครงสร้างการบริหารและคณะกรรมการ (รวมผู้แทนภาคธุรกิจ) "
    "• ก่อสร้าง/ปรับปรุงอาคาร AI & Robotics Innovation Lab "
    "• จัดหา GPU Computing Cluster และอุปกรณ์ "
    "• สรรหาบุคลากรหลัก (วิศวกร AI 5 คน, ผู้จัดการนวัตกรรม 2 คน) "
    "• เปิด MFU AI Bootcamp รุ่นแรก "
    "• เปิดรับ Startup Cohort แรก (5 ทีม) "
    "• ลงนาม MOU กับภาคเอกชน 3 ราย")

set_cell_text(t6.rows[2].cells[2],
    "• เปิดหลักสูตรปริญญาโท AI Engineering (รับ 30 คน/ปี) "
    "• เปิด Startup Incubation & Acceleration Program เต็มรูปแบบ (15 ทีม/รุ่น) "
    "• จัดตั้ง AI Solutions Consulting Unit ให้บริการ SME "
    "• สร้าง Regional AI Business Network ใน GMS 3 ประเทศ "
    "• จัด Demo Day / Investor Matching Event รายปี "
    "• พัฒนา AI Solutions สำหรับเกษตร สุขภาพ Smart City อย่างน้อย 10 โครงการ "
    "• ขยายพันธมิตรภาคเอกชนเป็น 20 ราย")

set_cell_text(t6.rows[3].cells[2],
    "• บรรลุ Self-sufficiency ทางการเงินบางส่วน (≥40%) จากรายได้ธุรกิจนวัตกรรม "
    "• Startup ที่บ่มเพาะสะสม 50 ทีม สร้างงาน 200+ ตำแหน่ง "
    "• ขยายเครือข่าย GMS ครบ 5 ประเทศ "
    "• เปิด Online AI Learning Platform "
    "• จัดตั้ง AI Technology Transfer & Licensing Office "
    "• ประเมินผลโครงการและวางแผนขยายระยะต่อไป")

# ─── TABLE 7: Budget — rebalance toward innovation ───────────────
print("  Table 7: Budget (rebalance)")
t7 = doc.tables[7]

# R9: Rename and increase innovation budget (take from research)
set_cell_text(t7.rows[9].cells[0], "   2.3 โครงการนวัตกรรมและบ่มเพาะธุรกิจ")
set_cell_text(t7.rows[9].cells[1], "3")
set_cell_text(t7.rows[9].cells[2], "12")
set_cell_text(t7.rows[9].cells[3], "5")
set_cell_text(t7.rows[9].cells[4], "20")

# R12: Increase hackathon/innovation day budget (take remaining from research)
set_cell_text(t7.rows[12].cells[0], "2.6 ค่าจัด Hackathon/Demo Day/Investor Event")
set_cell_text(t7.rows[12].cells[1], "2")
set_cell_text(t7.rows[12].cells[2], "5")
set_cell_text(t7.rows[12].cells[3], "3")
set_cell_text(t7.rows[12].cells[4], "10")
# Total adjustment: R9 was 25→20, R12 was 5→10, net same 30

# ─── TABLE 8: Annual budget — update descriptions ────────────────
print("  Table 8: Annual budget descriptions")
t8 = doc.tables[8]
set_cell_text(t8.rows[4].cells[2], "งบดำเนินการเป็นหลัก (80%) + งบลงทุนปรับปรุง (20%)")
set_cell_text(t8.rows[4].cells[3], "มุ่งนวัตกรรมและบ่มเพาะ Startup")
set_cell_text(t8.rows[5].cells[3], "สร้าง Self-sufficiency จากรายได้ธุรกิจนวัตกรรม")

# ─── TABLE 9: Sustainability plan — increase business revenue ────
print("  Table 9: Sustainability plan")
set_note_text(doc.tables[9].rows[0].cells[0],
    "แผนความยั่งยืน (หลังปีที่ 3 เป็นต้นไป): |"
    "ศูนย์ฯ จะสร้างรายได้จากแหล่งต่าง ๆ ได้แก่ "
    "(1) ค่าบริการฝึกอบรม/หลักสูตรระยะสั้น เป้าหมาย 10–15 ล้านบาท/ปี "
    "(2) รายได้จากการให้บริการ AI Solutions Consulting แก่ SME และภาคอุตสาหกรรม เป้าหมาย 15 ล้านบาท/ปี "
    "(3) ค่า Licensing Fees จากทรัพย์สินทางปัญญาและ Software License เป้าหมาย 10 ล้านบาท/ปี "
    "(4) Industry Partnership & Co-development Program เป้าหมาย 15 ล้านบาท/ปี "
    "และ (5) รายได้จาก Startup Equity/Revenue Sharing เป้าหมาย 10 ล้านบาท/ปี "
    "รวมรายได้เป้าหมายในปีที่ 5 ไม่น้อยกว่า 60 ล้านบาท/ปี หรือคิดเป็นร้อยละ 55 ของงบดำเนินการ")

# ─── TABLE 10: ROI analysis — update ─────────────────────────────
print("  Table 10: ROI analysis")
t10 = doc.tables[10]
set_cell_text(t10.rows[1].cells[1],
    "บวก – เมื่อคิดที่อัตราคิดลดร้อยละ 5 NPV อยู่ในระดับบวก เนื่องจากมูลค่าที่เกิดจาก Startup และนวัตกรรมเชิงพาณิชย์สูงกว่าต้นทุนลงทุน")
set_cell_text(t10.rows[2].cells[1],
    "7:1 – ทุก 1 บาทที่ลงทุน คาดสร้างมูลค่าทางเศรษฐกิจคืน 7 บาทภายใน 10 ปี ผ่าน Startup, AI Solutions และการจ้างงานมูลค่าสูง")
set_cell_text(t10.rows[3].cells[1],
    "• ต้นทุนฝึกอบรม 1 คน ≈ 15,000 บาท (เทียบเท่าตลาดเอกชน 50,000–80,000 บาท) "
    "• ต้นทุนบ่มเพาะ Startup 1 ทีม ≈ 500,000 บาท (Startup สร้างมูลค่าเฉลี่ย 5–10 ล้านบาท/ทีม)")
set_cell_text(t10.rows[4].cells[1],
    "ตั้งแต่ปีที่ 3 เป็นต้นไป ศูนย์ฯ สร้างรายได้ไม่น้อยกว่า 30 ล้านบาท/ปี จาก AI Consulting, Licensing, Industry Partnership และ Startup Revenue Sharing สู่เป้า 60 ล้านบาท/ปี ในปีที่ 5")

# ─── TABLE 12: Governance — add industry/startup rep ──────────────
print("  Table 12: Governance")
t12 = doc.tables[12]
set_cell_text(t12.rows[1].cells[1],
    "ประธาน: อธิการบดี มฟล. | สมาชิก: รองอธิการบดี, ผู้แทนกระทรวงดิจิทัลฯ, "
    "ผู้แทนภาคเอกชน/สภาอุตสาหกรรม 3 คน, ผู้แทน Startup Ecosystem 1 คน, "
    "ผู้ทรงคุณวุฒิภายนอก 2 คน | หน้าที่: กำหนดนโยบาย อนุมัติงบประมาณ ประเมินผลรายปี")
set_cell_text(t12.rows[2].cells[0], "คณะที่ปรึกษาด้านนวัตกรรมและธุรกิจ (Innovation Advisory Board)")
set_cell_text(t12.rows[2].cells[1],
    "ประกอบด้วยผู้เชี่ยวชาญ AI, ผู้ประกอบการ Startup สำเร็จ, นักลงทุน VC และผู้บริหารภาคอุตสาหกรรม 5–7 คน | "
    "หน้าที่: ให้คำแนะนำทิศทางนวัตกรรมและธุรกิจ, ประเมินศักยภาพ Startup, เชื่อมเครือข่ายนักลงทุนและอุตสาหกรรม")
set_cell_text(t12.rows[3].cells[1],
    "คุณสมบัติ: ดร./ผศ. ขึ้นไป มีประสบการณ์ AI ทั้งด้านวิจัยและอุตสาหกรรม/ธุรกิจ | "
    "หน้าที่: บริหารงานประจำ กำกับการดำเนินงาน รายงานต่อกรรมการ")
set_cell_text(t12.rows[4].cells[1],
    "Innovation & Business | Education | Applied Research | International Relations | แต่ละฝ่ายมีหัวหน้าฝ่ายรับผิดชอบ")

# ─── TABLE 14: Final proposal box ────────────────────────────────
print("  Table 14: Final proposal statement")
set_note_text(doc.tables[14].rows[0].cells[0],
    "ข้อเสนอขั้นสุดท้ายต่อผู้มีอำนาจตัดสินใจ|"
    "โครงการ MFU-ARIC เป็นการลงทุนเชิงยุทธศาสตร์ที่ตอบสนองต่อวาระแห่งชาติด้าน AI และดิจิทัล "
    "โดยมุ่งเน้นการสร้างระบบนิเวศนวัตกรรมและบ่มเพาะ Startup ด้าน AI ในภาคเหนือและอนุภูมิภาคลุ่มน้ำโขง "
    "มหาวิทยาลัยแม่ฟ้าหลวงมีความพร้อมทั้งวิชาการ โครงสร้างนวัตกรรม และเจตนารมณ์ "
    "การสนับสนุนจากภาครัฐด้วยงบประมาณ 150 ล้านบาทใน 5 ปี จะเป็นตัวเร่งสำคัญให้โครงการสร้าง Startup 50 ทีม "
    "สร้างงานมูลค่าสูง 200+ ตำแหน่ง และสร้างผลตอบแทนทางเศรษฐกิจไม่น้อยกว่า 2,000 ล้านบาท|"
    "มหาวิทยาลัยแม่ฟ้าหลวงจึงขอให้พิจารณาให้การสนับสนุนโครงการนี้ในทุกมิติ ทั้งงบประมาณ นโยบาย "
    "และการส่งเสริมเครือข่ายความร่วมมือ เพื่อให้ประเทศไทยสามารถก้าวสู่การเป็น AI Nation ที่ขับเคลื่อนด้วยนวัตกรรมและผู้ประกอบการในภูมิภาคอาเซียน")

# ─── TABLE 16: Milestones — add startup milestones ────────────────
print("  Table 16: Milestones")
t16 = doc.tables[16]
set_cell_text(t16.rows[2].cells[1],
    "เริ่มก่อสร้าง/ปรับปรุงอาคาร, จัดหา GPU ชุดแรก, เปิด MFU AI Bootcamp, รับ Startup Cohort แรก 5 ทีม, ลงนาม MOU เอกชน 3 ราย")
set_cell_text(t16.rows[3].cells[1],
    "เปิด Robotics Lab, เปิด AI Solutions Consulting Unit, เริ่มโครงการร่วมพัฒนากับ SME 5 โครงการ, จัด Demo Day ครั้งแรก")
set_cell_text(t16.rows[3].cells[2], "ฝ่ายนวัตกรรมและธุรกิจ")
set_cell_text(t16.rows[4].cells[1],
    "ลงนาม MOU ภาคเอกชน 10 ราย, เปิดหลักสูตร ป.โท AI รุ่นแรก, รับ Startup Cohort ที่ 2 (15 ทีม), Investor Matching Event")
set_cell_text(t16.rows[5].cells[1],
    "ขยาย Startup Incubation เป็น 30+ ทีม สะสม, จัดตั้ง AI Technology Transfer Office, สร้าง Regional AI Business Network, บรรลุรายได้ 30+ ล้าน/ปี")
set_cell_text(t16.rows[6].cells[1],
    "Startup สะสม 50 ทีม, บรรลุ KPI ปีที่ 5, รายได้ 60 ล้าน/ปี, ประเมินผลโครงการ, วางแผนขยายระยะ 2 (2575–2579)")

# ─── PARAGRAPHS ───────────────────────────────────────────────────
print("  Paragraphs: Objectives")

# P57: Shift from pure research to applied innovation
set_para_text(doc.paragraphs[57],
    "เพื่อสร้างนวัตกรรมและ AI Solutions ที่ตอบโจทย์ภาคอุตสาหกรรมและ SME ในภาคเหนือและอนุภูมิภาค "
    "รวมถึงบ่มเพาะ Startup ด้าน AI และ Robotics ให้เติบโตเป็นธุรกิจที่สร้างมูลค่าทางเศรษฐกิจ")

# P58: Strengthen innovation ecosystem language
set_para_text(doc.paragraphs[58],
    "เพื่อสร้างระบบนิเวศนวัตกรรมและผู้ประกอบการ (Innovation & Entrepreneurship Ecosystem) "
    "ที่เชื่อมโยงมหาวิทยาลัย ภาคอุตสาหกรรม นักลงทุน ภาครัฐ และชุมชน ขับเคลื่อนเศรษฐกิจดิจิทัลภาคเหนือ")

print("  Paragraphs: Funding sources")

# P82: Shift from research to innovation infrastructure
set_para_text(doc.paragraphs[82],
    "งบประมาณแผ่นดิน (ผ่านกระทรวงดิจิทัลฯ/สำนักงบประมาณ): 150 ล้านบาท (ร้อยละ 50) — "
    "สำหรับโครงสร้างพื้นฐานนวัตกรรม, บุคลากร, และการบ่มเพาะ Startup")

# P84: From research grants to innovation/business grants
set_para_text(doc.paragraphs[84],
    "ทุนนวัตกรรมจาก บพข./NIA/DEPA/สวทช.: 50 ล้านบาท (ร้อยละ 17) — "
    "สำหรับโครงการพัฒนานวัตกรรมเชิงพาณิชย์และ Startup Incubation ที่ขอทุนแบบ Competitive Grant")

# P85: From industry-sponsored research to co-development
set_para_text(doc.paragraphs[85],
    "Matching Fund ภาคเอกชน: 30 ล้านบาท (ร้อยละ 10) — "
    "จากพันธมิตรอุตสาหกรรม ผ่านรูปแบบ Industry Co-development และ Corporate Venture Partnership")

print("  Paragraphs: Economic benefits")

# P93: Increase economic impact projection
set_para_text(doc.paragraphs[93],
    "สร้างมูลค่าทางเศรษฐกิจไม่น้อยกว่า 2,000 ล้านบาทใน 5 ปี ผ่าน AI Startup, นวัตกรรมเชิงพาณิชย์ "
    "และการลงทุนต่อเนื่องของภาคเอกชน")

# P95: Add business/SME angle
set_para_text(doc.paragraphs[95],
    "เพิ่มขีดความสามารถในการแข่งขันของ SME และภาคอุตสาหกรรมในภาคเหนือ "
    "ด้วย AI Solutions สำหรับเกษตร อาหาร สุขภาพ การท่องเที่ยว และโลจิสติกส์")

# P96: Add startup job creation
set_para_text(doc.paragraphs[96],
    "สร้างงานที่มีมูลค่าสูงในภาคเหนือ ผ่าน Startup และธุรกิจนวัตกรรม AI "
    "ลดการย้ายถิ่นของแรงงานมีทักษะสู่กรุงเทพฯ เป้าหมาย 200+ ตำแหน่งใน 5 ปี")

# P106: Add entrepreneur attraction
set_para_text(doc.paragraphs[106],
    "ดึงดูดผู้ประกอบการ นักลงทุน และผู้เชี่ยวชาญด้าน AI ต่างชาติมาสู่ประเทศไทย "
    "สร้างรายได้จาก AI Startup Ecosystem และการศึกษา")

# ─── TABLE 11: Risk — add startup-related risk ───────────────────
print("  Table 11: Risk table (update R4)")
t11 = doc.tables[11]
set_cell_text(t11.rows[4].cells[0], "R4: Startup ล้มเหลวหรือขาดตลาด")
set_cell_text(t11.rows[4].cells[3],
    "คัดเลือก Startup อย่างเข้มงวดผ่าน Selection Committee, สร้าง Mentorship Network จากภาคเอกชน, "
    "จัด Investor Matching Event ทุกปี, ติดตาม Startup KPI รายไตรมาส")

# ─── Save funding proposal ───────────────────────────────────────
out1 = 'MFU_ARIC_Government_Funding_Proposal_2569.docx'
doc.save(out1)
print(f"\nSaved: {out1}")


# ═══════════════════════════════════════════════════════════════════
# PART 2: MFU_ARIC_Equipment_Specifications_Budget.docx
# ═══════════════════════════════════════════════════════════════════
print("\nLoading Equipment Budget...")
doc2 = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')

# ─── TABLE 3: Network/Cybersecurity rationale ────────────────────
print("  Table 3: Network rationale (add business angle)")
set_note_text(doc2.tables[3].rows[0].cells[0],
    "เหตุผลและความจำเป็น|"
    "ศูนย์ MFU-ARIC ต้องการระบบเครือข่ายที่มีความเร็วสูง (High-Performance Network) "
    "เพื่อรองรับการส่งข้อมูลขนาดใหญ่ระหว่าง GPU Cluster, Storage และ Workstation "
    "รวมถึงรองรับ AI Solutions Consulting Service และ Startup ที่ต้องเข้าถึง Computing Resources ได้อย่างปลอดภัย "
    "ระบบความปลอดภัยไซเบอร์ที่แข็งแกร่งจำเป็นสำหรับปกป้องทรัพย์สินทางปัญญา ข้อมูลภาคอุตสาหกรรม "
    "และข้อมูลธุรกิจของ Startup ที่เข้าร่วมโครงการ")

# ─── TABLE 10: GPU Cluster design philosophy ─────────────────────
print("  Table 10: GPU Cluster design (add startup/business use)")
set_note_text(doc2.tables[10].rows[0].cells[0],
    "หลักการออกแบบ GPU Cluster สำหรับศูนย์ MFU-ARIC|"
    "ออกแบบเป็น Hybrid Architecture ที่ผสมผสานระหว่าง On-premise High-Performance GPU Cluster "
    "สำหรับการพัฒนา AI Solutions และงานวิจัยประยุกต์ที่ต้องการประสิทธิภาพสูงและความปลอดภัยของข้อมูล "
    "กับ Cloud GPU Burst Capacity สำหรับ Startup และโครงการร่วมภาคอุตสาหกรรมที่ต้องการขยาย Scale เป็นครั้งคราว "
    "ซึ่งช่วยลดต้นทุนการลงทุนเริ่มต้นและเพิ่มความยืดหยุ่นในการรองรับธุรกิจนวัตกรรม|"
    "On-premise: เหมาะสำหรับการพัฒนา AI Product/Solution ที่ต้องการ Data Security, Low Latency และ Continuous Training|"
    "Cloud Hybrid: ใช้สำหรับ Startup Workload, Large-scale Experiment, Temporary Burst และ Backup DR|"
    "ผ่าน InfiniBand 200Gbps Interconnect เพื่อให้ GPU-to-GPU Communication มีความหน่วงต่ำที่สุด|"
    "Multi-tenant Architecture: รองรับการแบ่ง Resource ให้ Startup หลายทีมใช้งานพร้อมกันอย่างปลอดภัย")

# ─── TABLE 19: Robotics Lab design ───────────────────────────────
print("  Table 19: Robotics Lab design (add innovation/demo angle)")
set_note_text(doc2.tables[19].rows[0].cells[0],
    "แนวทางการออกแบบ Robotics Lab สำหรับ MFU-ARIC|"
    "ออกแบบเป็น Multi-purpose Robotics Innovation Space ที่รองรับการพัฒนานวัตกรรม การบ่มเพาะ Startup "
    "การเรียนการสอน และการสาธิตผลิตภัณฑ์ต่อนักลงทุนและภาคอุตสาหกรรม "
    "โดยแบ่งเป็น 4 โซนหลัก ได้แก่ "
    "(1) Collaborative Robotics Zone สำหรับพัฒนา Cobot Solutions ร่วมกับภาคอุตสาหกรรม "
    "(2) Autonomous Systems Zone สำหรับ Mobile Robot และ Drone Startup "
    "(3) AI-integrated Robotics Zone สำหรับพัฒนา AI+Robot Product/Solution ขั้นสูง "
    "และ (4) Demo & Investor Showcase Area สำหรับนำเสนอผลงาน Startup และผลิตภัณฑ์นวัตกรรมต่อนักลงทุนและภาคอุตสาหกรรม")

# ─── Save equipment budget ────────────────────────────────────────
out2 = 'MFU_ARIC_Equipment_Specifications_Budget.docx'
doc2.save(out2)
print(f"\nSaved: {out2}")

print("\n✓ Both documents updated successfully!")
print("  Key changes:")
print("  • KPIs: Publications 60→30, Patents 15→30, Startups 20→50, Revenue 40→60M")
print("  • Innovation Hub elevated to primary position in 4-Hub structure")
print("  • Work plan refocused on Startup Incubation & AI Solutions milestones")
print("  • Budget: Research 25M→20M, Hackathon/Events 5M→10M, relabeled as Innovation")
print("  • Governance: Added Startup/VC representation, renamed Advisory Board")
print("  • ROI raised to 7:1, economic impact 1,500M→2,000M")
print("  • Equipment design notes updated for Startup/business use cases")
