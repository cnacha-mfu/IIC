"""
fix_remaining.py — Fix stale references after budget enlargement (130→150, 300→320)

Fixes:
  1. P83: Gov channel อว., percentage 50→47
  2. P84: MFU co-funding 50→60M, percentage 17→19
  3. P85: Innovation funds 50→60M, percentage 17→19
  4. P86: Private percentage 10→9
  5. P87: International percentage 6→6 (unchanged, but total context changes)
  6. P91: Equipment 130→150, co-funding 150→170, breakdown update
  7. P151: MFU commitment 50→53%, 150→170M
  8. T7 R1/R6: Add CapEx/OpEx subtotals in section headers
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy


def set_cell_text(cell, text):
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')


def set_para_text(para, text):
    """Replace paragraph text preserving first run's formatting."""
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')


print("Loading Funding Proposal...")
doc = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# ─────────────────────────────────────────────────────────────
# FIX 1: P83 — Gov funding source
# ─────────────────────────────────────────────────────────────
print("\n[1] P83: Gov funding source")
set_para_text(doc.paragraphs[83],
    "งบประมาณแผ่นดิน (ผ่าน อว./สำนักงบประมาณ): 150 ล้านบาท (ร้อยละ 47) "
    "— สำหรับโครงสร้างพื้นฐานนวัตกรรม, ครุภัณฑ์ AI/Robotics, บุคลากร, "
    "และการบ่มเพาะ Startup")
print("     50→47%, กระทรวงดิจิทัลฯ→อว.")

# ─────────────────────────────────────────────────────────────
# FIX 2: P84 — MFU co-funding (50M→60M)
# ─────────────────────────────────────────────────────────────
print("\n[2] P84: MFU co-funding")
set_para_text(doc.paragraphs[84],
    "งบประมาณมหาวิทยาลัยแม่ฟ้าหลวง: 60 ล้านบาท (ร้อยละ 19) "
    "— แสดงความมุ่งมั่นและการร่วมลงทุนของสถาบัน "
    "สำหรับก่อสร้าง/ปรับปรุงอาคารและค่าดำเนินการเริ่มต้น")
print("     50→60M, 17→19%")

# ─────────────────────────────────────────────────────────────
# FIX 3: P85 — Innovation funds (50M→60M)
# ─────────────────────────────────────────────────────────────
print("\n[3] P85: Innovation funds")
set_para_text(doc.paragraphs[85],
    "ทุนนวัตกรรมจาก บพข./NIA/DEPA/สวทช.: 60 ล้านบาท (ร้อยละ 19) "
    "— สำหรับโครงการพัฒนานวัตกรรมเชิงพาณิชย์และ Startup Incubation "
    "ที่ขอทุนแบบ Competitive Grant")
print("     50→60M, 17→19%")

# ─────────────────────────────────────────────────────────────
# FIX 4: P86 — Private matching fund
# ─────────────────────────────────────────────────────────────
print("\n[4] P86: Private matching fund")
set_para_text(doc.paragraphs[86],
    "Matching Fund ภาคเอกชน: 30 ล้านบาท (ร้อยละ 9) "
    "— จากพันธมิตรอุตสาหกรรม ผ่านรูปแบบ Industry Co-development "
    "และ Corporate Venture Partnership")
print("     10→9%")

# ─────────────────────────────────────────────────────────────
# FIX 5: P87 — International funds
# ─────────────────────────────────────────────────────────────
print("\n[5] P87: International funds")
set_para_text(doc.paragraphs[87],
    "ทุนต่างประเทศ (EU, JICA, ADB, World Bank): 20 ล้านบาท (ร้อยละ 6) "
    "— สำหรับโครงการ GMS Cooperation และ Cross-border AI Initiatives")
print("     6% (unchanged)")

# ─────────────────────────────────────────────────────────────
# FIX 6: P91 — Budget note (equipment 130→150, co-funding 150→170)
# ─────────────────────────────────────────────────────────────
print("\n[6] P91: Budget note")
set_para_text(doc.paragraphs[91],
    "หมายเหตุงบประมาณ: "
    "(1) แผนจัดซื้อจัดจ้างครุภัณฑ์รวม 150 ล้านบาท "
    "จะจัดทำแผนจัดซื้อจัดจ้างประจำปีตาม "
    "พ.ร.บ.การจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 "
    "โดยเปิดเผยราคากลางและข้อกำหนดครุภัณฑ์ตามเอกสารแนบท้าย "
    "(2) งบสมทบจากมหาวิทยาลัยและแหล่งทุนอื่นรวม 170 ล้านบาท "
    "ประกอบด้วย งบประมาณมหาวิทยาลัย 60 ล้าน (มติสภามหาวิทยาลัย), "
    "กองทุน/ทุนวิจัยนวัตกรรม 60 ล้าน, "
    "Matching Fund ภาคเอกชน 30 ล้าน, "
    "ทุนต่างประเทศ 20 ล้าน — "
    "จะจัดทำ MOU กับแหล่งทุนทุกรายก่อนเริ่มโครงการ")
print("     130→150M equipment, 150→170M co-funding")

# ─────────────────────────────────────────────────────────────
# FIX 7: P151 — MFU commitment
# ─────────────────────────────────────────────────────────────
print("\n[7] P151: MFU commitment")
set_para_text(doc.paragraphs[151],
    "มหาวิทยาลัยแม่ฟ้าหลวงจะร่วมลงทุนไม่น้อยกว่าร้อยละ 53 ของวงเงินโครงการ "
    "(170 ล้านบาท) จากงบประมาณมหาวิทยาลัยและทุนวิจัยอื่น")
print("     50→53%, 150→170M")

# ─────────────────────────────────────────────────────────────
# FIX 8: T7 R1/R6 — Add CapEx/OpEx subtotals
# ─────────────────────────────────────────────────────────────
print("\n[8] T7: Add CapEx/OpEx subtotals in headers")
t7 = doc.tables[7]
set_cell_text(t7.rows[1].cells[0],
    "1. งบลงทุน (Capital Expenditure) — รวม 185 ล้านบาท")
set_cell_text(t7.rows[6].cells[0],
    "2. งบดำเนินการ (Operating Expenditure) — รวม 135 ล้านบาท")
print("     R1: CapEx total 185M, R6: OpEx total 135M")

# Save
doc.save('MFU_ARIC_Government_Funding_Proposal_2569.docx')
print("\n✓ Saved.")

# ─────────────────────────────────────────────────────────────
# VERIFICATION
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("VERIFICATION")
print("=" * 60)

doc2 = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

print("\n--- Funding Sources (P83-87) ---")
for pi in range(83, 88):
    t = doc2.paragraphs[pi].text.strip()
    if t:
        print(f"  P{pi}: {t[:100]}")

print("\n--- Budget Note (P91) ---")
t91 = doc2.paragraphs[91].text.strip()
print(f"  P91: {t91[:120]}")
# Check amounts
assert '150 ล้านบาท' in t91, "P91 should say 150M equipment"
assert '170 ล้านบาท' in t91, "P91 should say 170M co-funding"

print("\n--- MFU Commitment (P151) ---")
print(f"  P151: {doc2.paragraphs[151].text.strip()[:120]}")

print("\n--- T7 Headers ---")
print(f"  R1: {doc2.tables[7].rows[1].cells[0].text.strip()[:60]}")
print(f"  R6: {doc2.tables[7].rows[6].cells[0].text.strip()[:60]}")

# Cross-check funding sources sum
print("\n--- Funding Sum Check ---")
sources = {
    'Gov': 150, 'MFU': 60, 'Innovation': 60,
    'Private': 30, 'International': 20
}
total = sum(sources.values())
print(f"  Sources: {sources}")
print(f"  Sum: {total}M (expect 320) {'✓' if total==320 else '✗'}")

# CapEx/OpEx check
capex = 35 + 67 + 68 + 15
opex = 65 + 20 + 20 + 10 + 10 + 10
print(f"\n  CapEx: {capex}M (expect 185) {'✓' if capex==185 else '✗'}")
print(f"  OpEx: {opex}M (expect 135) {'✓' if opex==135 else '✗'}")
print(f"  Total: {capex+opex}M (expect 320) {'✓' if capex+opex==320 else '✗'}")

print("\n" + "=" * 60)
print("ALL REMAINING FIXES DONE")
print("=" * 60)
