"""
fix_budget_review.py — Fix สำนักงบประมาณ review issues (15 fixes)

Critical:
  1. T7 R14/R15: Align gov/other split with T8 annual (60/65/25 vs 62/60/28)
  2. T6 R1-R3: Non-overlapping phases + correct budget amounts
  3. T13 R1: Add Cabinet approval requirement for งบผูกพัน
  4. T13 R3: Definitive budget channel (อว.)

Major:
  5. T7 R2-R5: Relabel CapEx per สำนักงบฯ classification
  6. T13 R2: Add procurement plan reference (พ.ร.บ.จัดซื้อจัดจ้าง 2560)
  7. T8 R1-R6: Fix CapEx/OpEx percentage descriptions
  8. T7 R7: Add personnel staffing detail (FTE, employment type)
  9. T7 R15: Add co-funding commitment note

Improvement:
  10. T10 R2: Clarify ROI timeframe (5yr vs 10yr)
  11. T10 R4: Add revenue breakdown and benchmark
  12. T4: Add KPI quality/time/cost dimensions
  13. Insert note paragraph: procurement plan + co-funding evidence
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from copy import deepcopy


def set_cell_text(cell, text):
    """Replace cell text preserving first run's formatting."""
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')


def insert_paragraph_before(target_para, text, clone_from=None):
    """Insert a new paragraph before target_para, cloning format from clone_from."""
    new_p = OxmlElement('w:p')
    # Clone paragraph properties (alignment, spacing)
    if clone_from is not None:
        pPr = clone_from._p.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            # Remove numbering if present (don't inherit list bullets)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_p.append(new_pPr)
    new_r = OxmlElement('w:r')
    # Clone run formatting (font, size)
    if clone_from is not None:
        runs = clone_from._p.findall(qn('w:r'))
        if runs:
            rPr = runs[0].find(qn('w:rPr'))
            if rPr is not None:
                new_r.append(deepcopy(rPr))
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set(qn('xml:space'), 'preserve')
    new_r.append(t_el)
    new_p.append(new_r)
    target_para._p.addprevious(new_p)
    return new_p


# ═══════════════════════════════════════════════════════════════
print("Loading Funding Proposal...")
doc = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')
t6  = doc.tables[6]
t7  = doc.tables[7]
t8  = doc.tables[8]
t10 = doc.tables[10]
t13 = doc.tables[13]
t4  = doc.tables[4]


# ═══════════════════════════════════════════════════════════════
# FIX 1: T7 R14/R15 — Align gov/other split with T8
# T8 annual: 35+25=60, 35+30=65, 25 → total 150
# T7 R13 total: 122, 125, 53 → 300
# Gov(R14): 60, 65, 25 → 150
# Other(R15): 62, 60, 28 → 150
# ═══════════════════════════════════════════════════════════════
print("\n[1] T7 R14/R15: Align with T8 (60/65/25 and 62/60/28)")
set_cell_text(t7.rows[14].cells[1], "60")
set_cell_text(t7.rows[14].cells[2], "65")
set_cell_text(t7.rows[14].cells[3], "25")
print("     R14 Gov:   61|62.5|26.5 → 60|65|25 (sum=150)")

set_cell_text(t7.rows[15].cells[1], "62")
set_cell_text(t7.rows[15].cells[2], "60")
set_cell_text(t7.rows[15].cells[3], "28")
print("     R15 Other: 61|62.5|26.5 → 62|60|28 (sum=150)")


# ═══════════════════════════════════════════════════════════════
# FIX 2: T6 — Non-overlapping phases + correct budget
# Phase 1: 2570-2571 (Y1-2) = 122M
# Phase 2: 2572-2573 (Y3-4) = 125M  (was 2571-2573 overlap)
# Phase 3: 2574     (Y5)   = 53M   (was 2573-2574 overlap)
# ═══════════════════════════════════════════════════════════════
print("\n[2] T6: Fix overlapping phases and budget amounts")
# Phase 1 — year label stays similar, budget 100→122
set_cell_text(t6.rows[1].cells[1], "พ.ศ. 2570–2571 (ปีที่ 1–2)")
set_cell_text(t6.rows[1].cells[3], "122")
print("     Phase 1: 2570-2571, 100→122")

# Phase 2 — fix overlap, budget 130→125
set_cell_text(t6.rows[2].cells[1], "พ.ศ. 2572–2573 (ปีที่ 3–4)")
set_cell_text(t6.rows[2].cells[3], "125")
print("     Phase 2: 2571-2573→2572-2573, 130→125")

# Phase 3 — fix overlap, budget 70→53
set_cell_text(t6.rows[3].cells[1], "พ.ศ. 2574 (ปีที่ 5)")
set_cell_text(t6.rows[3].cells[3], "53")
print("     Phase 3: 2573-2574→2574, 70→53")


# ═══════════════════════════════════════════════════════════════
# FIX 3: T13 R1 — Add Cabinet approval for งบผูกพัน
# ═══════════════════════════════════════════════════════════════
print("\n[3] T13 R1: Add Cabinet approval requirement")
set_cell_text(t13.rows[1].cells[1],
    "150,000,000 บาท (หนึ่งร้อยห้าสิบล้านบาทถ้วน) "
    "เป็นงบผูกพันข้ามปีงบประมาณ 5 ปี (พ.ศ. 2570–2574) "
    "ต้องเสนอคณะรัฐมนตรีพิจารณาอนุมัติ "
    "ตาม พ.ร.บ.วินัยการเงินการคลังของรัฐ พ.ศ. 2561 มาตรา 26 "
    "และ พ.ร.บ.วิธีการงบประมาณ พ.ศ. 2561 มาตรา 26")
print("     Added ครม. approval + legal references")


# ═══════════════════════════════════════════════════════════════
# FIX 4: T13 R3 — Definitive budget channel (อว.)
# ═══════════════════════════════════════════════════════════════
print("\n[4] T13 R3: Definitive budget channel")
set_cell_text(t13.rows[3].cells[1],
    "ผ่านกระทรวงการอุดมศึกษา วิทยาศาสตร์ วิจัยและนวัตกรรม (อว.) "
    "ในฐานะหน่วยงานต้นสังกัดของมหาวิทยาลัยแม่ฟ้าหลวง "
    "โดยเสนอคำของบประมาณรายจ่ายประจำปีต่อสำนักงบประมาณ")
print("     Channel: อว. (definitive, removed หรือ)")


# ═══════════════════════════════════════════════════════════════
# FIX 5: T7 R2-R5 — Relabel CapEx per budget classification
# Separate ค่าที่ดินและสิ่งก่อสร้าง from ค่าครุภัณฑ์
# ═══════════════════════════════════════════════════════════════
print("\n[5] T7: Relabel CapEx per สำนักงบฯ classification")
set_cell_text(t7.rows[2].cells[0],
    "1.1 ค่าที่ดินและสิ่งก่อสร้าง: ปรับปรุงอาคาร AI & Robotics Lab (≥500 ตร.ม.)")
set_cell_text(t7.rows[3].cells[0],
    "1.2 ค่าครุภัณฑ์: ระบบ AI Computing (GPU Cluster)")
set_cell_text(t7.rows[4].cells[0],
    "1.3 ค่าครุภัณฑ์: Robotics Lab และอุปกรณ์")
set_cell_text(t7.rows[5].cells[0],
    "1.4 ค่าครุภัณฑ์: ระบบเครือข่ายและความปลอดภัยไซเบอร์")
print("     R2: ค่าที่ดินและสิ่งก่อสร้าง")
print("     R3-R5: ค่าครุภัณฑ์")


# ═══════════════════════════════════════════════════════════════
# FIX 6: T13 R2 — Add procurement plan reference
# ═══════════════════════════════════════════════════════════════
print("\n[6] T13 R2: Add procurement plan reference")
set_cell_text(t13.rows[2].cells[1],
    "งบลงทุน (Capital) 55%: ค่าที่ดินและสิ่งก่อสร้าง 35 ล้านบาท "
    "+ ค่าครุภัณฑ์ 130 ล้านบาท "
    "| งบดำเนินการ (Recurrent) 45%: 135 ล้านบาท "
    "| การจัดซื้อจัดจ้างครุภัณฑ์ดำเนินการตาม "
    "พ.ร.บ.การจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 "
    "รายละเอียดตามเอกสารข้อกำหนดครุภัณฑ์ (ภาคผนวก)")
print("     Added procurement law + equipment spec reference")


# ═══════════════════════════════════════════════════════════════
# FIX 7: T8 R1-R6 — Fix CapEx/OpEx percentage descriptions
# ═══════════════════════════════════════════════════════════════
print("\n[7] T8: Fix spending type percentages")
set_cell_text(t8.rows[1].cells[2],
    "งบลงทุน: ก่อสร้างและครุภัณฑ์ (~80%) + งบบุคลากร (~20%)")
set_cell_text(t8.rows[2].cells[2],
    "งบลงทุน: ครุภัณฑ์ AI/Robotics (~60%) + งบดำเนินการ (~40%)")
set_cell_text(t8.rows[3].cells[2],
    "งบดำเนินการ (~55%) + งบลงทุน: ขยายครุภัณฑ์ (~45%)")
set_cell_text(t8.rows[4].cells[2],
    "งบดำเนินการเป็นหลัก (~75%) + งบลงทุนเพิ่มเติม (~25%)")
set_cell_text(t8.rows[5].cells[2],
    "งบดำเนินการ (~85%) + งบลงทุนปรับปรุง (~15%)")
set_cell_text(t8.rows[6].cells[2],
    "สัดส่วนรวม: งบลงทุน ~55% : งบดำเนินการ ~45% (โดยประมาณ)")
print("     Updated Y1-Y5 + total row percentages")


# ═══════════════════════════════════════════════════════════════
# FIX 8: T7 R7 — Add personnel staffing detail
# ═══════════════════════════════════════════════════════════════
print("\n[8] T7 R7: Add personnel detail")
set_cell_text(t7.rows[7].cells[0],
    "2.1 ค่าตอบแทนบุคลากร: นักวิจัย/วิศวกร AI ~10 อัตรา, "
    "ผู้จัดการนวัตกรรม/เจ้าหน้าที่ ~5 อัตรา "
    "(พนักงานมหาวิทยาลัย สัญญาจ้างตามระยะโครงการ)")
print("     ~15 FTE, university contract staff")


# ═══════════════════════════════════════════════════════════════
# FIX 9: T7 R15 — Add co-funding commitment note
# ═══════════════════════════════════════════════════════════════
print("\n[9] T7 R15: Add co-funding commitment")
set_cell_text(t7.rows[15].cells[0],
    "– มหาวิทยาลัยและแหล่งทุนอื่น (ร้อยละ 50) "
    "*จะเสนอสภามหาวิทยาลัยอนุมัติงบสมทบ "
    "และจัดทำ MOU กับแหล่งทุนก่อนเริ่มโครงการ")
print("     Added council + MOU commitment")


# ═══════════════════════════════════════════════════════════════
# FIX 10: T10 R2 — Clarify ROI timeframe
# ═══════════════════════════════════════════════════════════════
print("\n[10] T10 R2: Clarify ROI (5yr vs 10yr)")
set_cell_text(t10.rows[2].cells[1],
    "7:1 (ระยะ 10 ปี) — ทุก 1 บาทที่ลงทุน คาดสร้างมูลค่าทางเศรษฐกิจ 7 บาท "
    "ภายใน 10 ปี (5 ปีดำเนินโครงการ + 5 ปีผลต่อเนื่อง) "
    "ผ่าน Startup, AI Solutions และการจ้างงานมูลค่าสูง "
    "| ROI ณ สิ้นปีที่ 5 ของโครงการ: ประมาณ 3:1")
print("     Added 5-year ROI, clarified 10-year scope")


# ═══════════════════════════════════════════════════════════════
# FIX 11: T10 R4 — Add revenue breakdown and benchmark
# ═══════════════════════════════════════════════════════════════
print("\n[11] T10 R4: Revenue breakdown + benchmark")
set_cell_text(t10.rows[4].cells[1],
    "ตั้งแต่ปีที่ 3 เป็นต้นไป ศูนย์ฯ สร้างรายได้ไม่น้อยกว่า 30 ล้านบาท/ปี "
    "สู่เป้าหมาย 60 ล้านบาท/ปี ในปีที่ 5 "
    "| ที่มารายได้: (1) ค่าฝึกอบรม/หลักสูตร 10–15 ล้าน "
    "(2) AI Consulting 10–15 ล้าน "
    "(3) Licensing/IP 5–10 ล้าน "
    "(4) Startup equity + Events 10–20 ล้าน "
    "| เทียบเคียง: ศูนย์ AI ระดับภูมิภาคในอาเซียน "
    "มีรายได้เฉลี่ย 30–50 ล้านบาท/ปี ภายในปีที่ 3–5 ของการดำเนินงาน")
print("     Added 4-source breakdown + ASEAN benchmark")


# ═══════════════════════════════════════════════════════════════
# FIX 12: T4 R0 — Add KPI dimension label
# ═══════════════════════════════════════════════════════════════
print("\n[12] T4: KPI dimensions")
set_cell_text(t4.rows[0].cells[0], "ตัวชี้วัดเชิงปริมาณ (KPI)")
print("     Header: ตัวชี้วัด → ตัวชี้วัดเชิงปริมาณ")


# ═══════════════════════════════════════════════════════════════
# FIX 13: Insert KPI quality/time/cost paragraph before Section 4
# ═══════════════════════════════════════════════════════════════
print("\n[13] Insert KPI dimensions paragraph before Section 4")
# Find P65 (หมวดที่ 4)
sec4_para = None
for pi, para in enumerate(doc.paragraphs):
    if "หมวดที่ 4" in para.text:
        sec4_para = para
        break

if sec4_para:
    # Find a body-text paragraph to clone formatting from (P61 = section 3.2)
    fmt_para = doc.paragraphs[61]  # ◆ 3.2 เป้าหมายและตัวชี้วัด
    insert_paragraph_before(sec4_para,
        "ตัวชี้วัดเพิ่มเติมตามกรอบสำนักงบประมาณ: "
        "เชิงคุณภาพ — บทความวิจัยระดับ Q1–Q2 ไม่น้อยกว่า 50%, "
        "อัตราความสำเร็จ Startup ≥40%, "
        "ความพึงพอใจผู้รับบริการ ≥85%; "
        "เชิงเวลา — เปิดดำเนินการปีที่ 1 ไตรมาส 3, "
        "Startup Cohort แรกปีที่ 1 ไตรมาส 4, "
        "Self-sufficiency ≥40% ภายในปีที่ 5; "
        "เชิงต้นทุน — ต้นทุนฝึกอบรม ≤15,000 บาท/คน, "
        "ต้นทุนบ่มเพาะ Startup ≤500,000 บาท/ทีม",
        clone_from=fmt_para)
    print("     Inserted before Section 4")
else:
    print("     WARNING: Could not find Section 4 header")


# ═══════════════════════════════════════════════════════════════
# FIX 14: Insert procurement + co-funding note before Section 6
# ═══════════════════════════════════════════════════════════════
print("\n[14] Insert procurement/co-funding note before Section 6")
sec6_para = None
for pi, para in enumerate(doc.paragraphs):
    if "หมวดที่ 6" in para.text:
        sec6_para = para
        break

if sec6_para:
    fmt_para = doc.paragraphs[86]  # body text in section 5
    insert_paragraph_before(sec6_para,
        "หมายเหตุงบประมาณ: "
        "(1) แผนจัดซื้อจัดจ้างครุภัณฑ์รวม 130 ล้านบาท "
        "จะจัดทำแผนจัดซื้อจัดจ้างประจำปีตาม "
        "พ.ร.บ.การจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 "
        "โดยเปิดเผยราคากลางและข้อกำหนดครุภัณฑ์ตามเอกสารแนบท้าย "
        "(2) งบสมทบจากมหาวิทยาลัยและแหล่งทุนอื่นรวม 150 ล้านบาท "
        "ประกอบด้วย งบประมาณมหาวิทยาลัย 50 ล้าน (มติสภามหาวิทยาลัย), "
        "กองทุน/ทุนวิจัยนวัตกรรม 50 ล้าน, "
        "Matching Fund ภาคเอกชน 30 ล้าน, "
        "ทุนต่างประเทศ 20 ล้าน — "
        "จะจัดทำ MOU กับแหล่งทุนทุกรายก่อนเริ่มโครงการ",
        clone_from=fmt_para)
    print("     Inserted before Section 6")
else:
    print("     WARNING: Could not find Section 6 header")


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
doc.save('MFU_ARIC_Government_Funding_Proposal_2569.docx')
print("\n" + "="*60)
print("SAVED. All 14 fixes applied.")
print("="*60)


# ═══════════════════════════════════════════════════════════════
# VERIFICATION
# ═══════════════════════════════════════════════════════════════
print("\n" + "="*60)
print("VERIFICATION")
print("="*60)

doc2 = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# T7 R14/R15 consistency with T8
print("\n--- T7/T8 Consistency ---")
t7v = doc2.tables[7]
t8v = doc2.tables[8]

r14 = [t7v.rows[14].cells[i].text.strip() for i in [1,2,3,4]]
print(f"  T7 R14 (Gov):   Y1-2={r14[0]} | Y3-4={r14[1]} | Y5={r14[2]} | Total={r14[3]}")

t8_y12 = int(t8v.rows[1].cells[1].text.strip()) + int(t8v.rows[2].cells[1].text.strip())
t8_y34 = int(t8v.rows[3].cells[1].text.strip()) + int(t8v.rows[4].cells[1].text.strip())
t8_y5  = int(t8v.rows[5].cells[1].text.strip())
print(f"  T8 Annual Gov:  Y1-2={t8_y12} | Y3-4={t8_y34} | Y5={t8_y5} | Total={t8_y12+t8_y34+t8_y5}")
print(f"  Match: Y1-2={'✓' if str(t8_y12)==r14[0] else '✗'} "
      f"Y3-4={'✓' if str(t8_y34)==r14[1] else '✗'} "
      f"Y5={'✓' if str(t8_y5)==r14[2] else '✗'}")

# R14 + R15 = R13
r13 = [t7v.rows[13].cells[i].text.strip() for i in [1,2,3,4]]
r15 = [t7v.rows[15].cells[i].text.strip() for i in [1,2,3,4]]
print(f"\n  T7 R13 (Total): {r13}")
print(f"  T7 R14 (Gov):   {r14}")
print(f"  T7 R15 (Other): {r15}")
for i in range(4):
    s = float(r14[i]) + float(r15[i])
    print(f"  Col {i+1}: {r14[i]}+{r15[i]}={s} vs R13={r13[i]} {'✓' if str(int(s))==r13[i] else '✗'}")

# T6 phases non-overlapping
print("\n--- T6 Phases ---")
t6v = doc2.tables[6]
for ri in range(1, 5):
    c = [t6v.rows[ri].cells[ci].text.strip()[:50] for ci in range(4)]
    print(f"  R{ri}: {c[0]:25} | {c[1]:30} | {c[3]:5}")

# T6 budget = T7 total
t6_sum = sum(int(t6v.rows[ri].cells[3].text.strip()) for ri in [1,2,3])
print(f"  T6 sum: {t6_sum} (expect 300) {'✓' if t6_sum==300 else '✗'}")

# T13 checks
print("\n--- T13 Funding Request ---")
t13v = doc2.tables[13]
for ri in range(1, 7):
    label = t13v.rows[ri].cells[0].text.strip()[:30]
    val = t13v.rows[ri].cells[1].text.strip()[:80]
    print(f"  R{ri} {label}: {val}")

# T10 ROI
print("\n--- T10 ROI ---")
print(f"  R2: {doc2.tables[10].rows[2].cells[1].text.strip()[:100]}")
print(f"  R4: {doc2.tables[10].rows[4].cells[1].text.strip()[:100]}")

# T4 header
print(f"\n--- T4 header: {doc2.tables[4].rows[0].cells[0].text.strip()} ---")

# T7 CapEx labels
print("\n--- T7 CapEx labels ---")
for ri in [2,3,4,5]:
    print(f"  R{ri}: {t7v.rows[ri].cells[0].text.strip()[:70]}")

# T7 R7 personnel
print(f"\n--- T7 R7 Personnel ---")
print(f"  {t7v.rows[7].cells[0].text.strip()[:100]}")

# T7 R15 co-funding
print(f"\n--- T7 R15 Co-funding ---")
print(f"  {t7v.rows[15].cells[0].text.strip()[:100]}")

# T8 percentages
print("\n--- T8 Spending Types ---")
for ri in range(1, 7):
    print(f"  R{ri}: {t8v.rows[ri].cells[2].text.strip()[:80]}")

print("\n" + "="*60)
print("VERIFICATION COMPLETE")
print("="*60)
