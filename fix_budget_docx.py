"""
Fix all issues in MFU_ARIC_Equipment_Specifications_Budget.docx

Changes:
1. Fill in all missing quantities
2. Phase cybersecurity (defer PAM, NAC, DLP, Cyber Range to Phase 2)
3. Reduce storage to fit Category 1 = 15M
4. Reduce H100 from 2→1, add Cloud GPU Credits to fit Category 2 = 65M
5. Replace da Vinci simulator with MicroPort Toumai trainer
6. Add Humanoid Robot (Unitree G1) to Category 3
7. Fix all subtotals and totals to be arithmetically correct
8. Fix item counts in summary table
9. Align procurement phases with funding proposal
10. Add license renewal note
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell
from copy import deepcopy

INPUT_FILE = 'MFU_ARIC_Equipment_Specifications_Budget.docx'
OUTPUT_FILE = 'MFU_ARIC_Equipment_Specifications_Budget.docx'

doc = Document(INPUT_FILE)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_text(cell, text, bold=False, color='1A1A2E', size_val='32', alignment=None):
    """Set cell text with explicit formatting."""
    para = cell.paragraphs[0]
    if alignment is not None:
        align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                     'left': WD_ALIGN_PARAGRAPH.LEFT,
                     'right': WD_ALIGN_PARAGRAPH.RIGHT}
        para.alignment = align_map[alignment]
    # Remove all existing runs
    for r_elem in para._p.findall(qn('w:r')):
        para._p.remove(r_elem)
    # Create new run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), size_val)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), size_val)
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    para._p.append(r)


def remove_table_row(table, row_index):
    """Remove a row from a table by index."""
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)


def clone_row_after(table, source_row_index, insert_after_index):
    """Clone a row and insert after specified row. Returns new tr element."""
    source_tr = table.rows[source_row_index]._tr
    new_tr = deepcopy(source_tr)
    ref_tr = table.rows[insert_after_index]._tr
    ref_tr.addnext(new_tr)
    return new_tr


def set_tr_cell(tr_element, col_index, text, bold=False, color='1A1A2E',
                size_val='32', alignment=None):
    """Set text in a cell of a raw w:tr element."""
    tcs = tr_element.findall(qn('w:tc'))
    if col_index < len(tcs):
        cell = _Cell(tcs[col_index], None)
        set_cell_text(cell, text, bold=bold, color=color,
                     size_val=size_val, alignment=alignment)


def set_note_text(table, text):
    """Replace text in a single-cell note table."""
    cell = table.rows[0].cells[0]
    # Clear all paragraphs except first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._p
        cell._tc.remove(p)
    # Clear first paragraph runs
    para = cell.paragraphs[0]
    for r_elem in para._p.findall(qn('w:r')):
        para._p.remove(r_elem)
    # Split text by | into separate runs (mimicking original format)
    parts = text.split(' | ')
    for i, part in enumerate(parts):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), '1A1A2E')
        rPr.append(color_elem)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '28')
        rPr.append(sz)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        separator = ' | ' if i < len(parts) - 1 else ''
        t.text = part + separator
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        para._p.append(r)


def fmt(value):
    """Format number with commas."""
    return f"{value:,}"


# ============================================================
# TABLE INDEX REFERENCE (from extraction):
# 0  = Cover/metadata
# 1  = Pricing disclaimer
# 2  = Category 1 header
# 3  = Category 1 justification
# 4  = Network items (1.1-1.7)
# 5  = Cybersecurity items (2.1-2.10)
# 6  = Storage items (3.1-3.3)
# 7  = Category 1 total
# 8  = Category 1 notes
# 9  = Category 2 header
# 10 = Category 2 design philosophy
# 11 = GPU servers (4.1-4.3)
# 12 = GPU capability note
# 13 = CPU servers (5.1-5.2)
# 14 = Workstation & Software (6.1-6.4)
# 15 = DC Infrastructure (7.1-7.5)
# 16 = Category 2 total
# 17 = Category 2 notes
# 18 = Category 3 header
# 19 = Category 3 design philosophy
# 20 = Cobots (8.1-8.6)
# 21 = Autonomous systems (9.1-9.5)
# 22 = Sensing & Vision (10.1-10.6)
# 23 = Simulation & Software (11.1-11.4)
# 24 = Lab Equipment (12.1-12.7)
# 25 = Category 3 total
# 26 = Category 3 notes
# 27 = Grand summary note
# 28 = Summary table
# 29 = Procurement phase table
# 30 = Procurement notes
# ============================================================

print("=== Starting document fixes ===\n")

# ============================================================
# STEP 1: Fill quantities in ALL data tables
# ============================================================
print("Step 1: Filling in quantities...")

# Table 4: Network (1.1-1.7), rows 1-7
t4 = doc.tables[4]
qtys_4 = {1: '2', 2: '6', 3: '1', 4: '8', 5: '1', 6: '12', 7: '1'}
for ri, q in qtys_4.items():
    set_cell_text(t4.rows[ri].cells[3], q, alignment='center')

# Table 5: Cybersecurity (2.1-2.10), rows 1-10
t5 = doc.tables[5]
qtys_5 = {1: '2', 2: '1', 3: '1', 4: '1', 5: '1', 6: '1', 7: '4', 8: '4', 9: '1', 10: '1'}
for ri, q in qtys_5.items():
    set_cell_text(t5.rows[ri].cells[3], q, alignment='center')

# Table 6: Storage (3.1-3.3), rows 1-3
t6 = doc.tables[6]
qtys_6 = {1: '1', 2: '1', 3: '1'}
for ri, q in qtys_6.items():
    set_cell_text(t6.rows[ri].cells[3], q, alignment='center')

# Table 11: GPU (4.1-4.3), rows 1-3
t11 = doc.tables[11]
qtys_11 = {1: '2', 2: '4', 3: '2'}
for ri, q in qtys_11.items():
    set_cell_text(t11.rows[ri].cells[3], q, alignment='center')

# Table 13: CPU (5.1-5.2), rows 1-2
t13 = doc.tables[13]
qtys_13 = {1: '2', 2: '4'}
for ri, q in qtys_13.items():
    set_cell_text(t13.rows[ri].cells[3], q, alignment='center')

# Table 14: WS/SW (6.1-6.4), rows 1-4
t14 = doc.tables[14]
qtys_14 = {1: '10', 2: '20', 3: '1', 4: '1'}
for ri, q in qtys_14.items():
    set_cell_text(t14.rows[ri].cells[3], q, alignment='center')

# Table 15: DC Infra (7.1-7.5), rows 1-5
t15 = doc.tables[15]
qtys_15 = {1: '10', 2: '4', 3: '2', 4: '10', 5: '1'}
for ri, q in qtys_15.items():
    set_cell_text(t15.rows[ri].cells[3], q, alignment='center')

# Table 20: Cobots (8.1-8.6), rows 1-6
t20 = doc.tables[20]
qtys_20 = {1: '4', 2: '4', 3: '2', 4: '2', 5: '4', 6: '4'}
for ri, q in qtys_20.items():
    set_cell_text(t20.rows[ri].cells[3], q, alignment='center')

# Table 21: Autonomous (9.1-9.5), rows 1-5
t21 = doc.tables[21]
qtys_21 = {1: '3', 2: '6', 3: '1', 4: '4', 5: '1'}
for ri, q in qtys_21.items():
    set_cell_text(t21.rows[ri].cells[3], q, alignment='center')

# Table 22: Sensing (10.1-10.6), rows 1-6
t22 = doc.tables[22]
qtys_22 = {1: '1', 2: '4', 3: '20', 4: '8', 5: '10', 6: '4'}
for ri, q in qtys_22.items():
    set_cell_text(t22.rows[ri].cells[3], q, alignment='center')

# Table 23: Sim/SW (11.1-11.4), rows 1-4
t23 = doc.tables[23]
qtys_23 = {1: '1', 2: '1', 3: '1', 4: '5'}
for ri, q in qtys_23.items():
    set_cell_text(t23.rows[ri].cells[3], q, alignment='center')

# Table 24: Lab Equipment (12.1-12.7), rows 1-7
t24 = doc.tables[24]
qtys_24 = {1: '3', 2: '4', 3: '1', 4: '2', 5: '1', 6: '1', 7: '2'}
for ri, q in qtys_24.items():
    set_cell_text(t24.rows[ri].cells[3], q, alignment='center')

print("  Quantities filled in all 12 data tables.\n")

# ============================================================
# STEP 2: CATEGORY 1 - Phase cybersecurity, reduce storage
# Target: 15,000,000
# ============================================================
print("Step 2: Fixing Category 1 (target 15M)...")

# --- Table 5: Cybersecurity ---
# Remove rows for deferred items (bottom to top to preserve indices):
# Row 10 = item 2.10 Cyber Range
# Row 6  = item 2.6  DLP
# Row 5  = item 2.5  NAC
# Row 4  = item 2.4  PAM
for ri in [10, 6, 5, 4]:
    remove_table_row(t5, ri)

# After removal, table 5 now has 8 rows:
# Row 0: header
# Row 1: 2.1 Firewall (was row 1) - qty 2, 850K, 1.7M - keep
# Row 2: 2.2 SIEM (was row 2) - qty 1, 1.2M, 1.2M - keep
# Row 3: 2.3 EDR (was row 3) - qty 1, 450K, 450K - keep
# Row 4: 2.7 SOC WS (was row 7) - renumber to 2.4, reduce qty 4→2
# Row 5: 2.8 UPS (was row 8) - renumber to 2.5, reduce qty 4→2
# Row 6: 2.9 Vuln Scanner (was row 9) - renumber to 2.6
# Row 7: subtotal (was row 11)

# Re-read table after row removal
t5 = doc.tables[5]

# Renumber and adjust SOC Workstation (now row 4)
set_cell_text(t5.rows[4].cells[0], '2.4', alignment='center')
set_cell_text(t5.rows[4].cells[3], '2', alignment='center')  # qty 4→2
set_cell_text(t5.rows[4].cells[6], fmt(240_000), bold=True, color='1A3A6B',
              alignment='center')

# Renumber UPS (now row 5), reduce qty 4→2
set_cell_text(t5.rows[5].cells[0], '2.5', alignment='center')
set_cell_text(t5.rows[5].cells[3], '2', alignment='center')  # qty 4→2
set_cell_text(t5.rows[5].cells[6], fmt(170_000), bold=True, color='1A3A6B',
              alignment='center')

# Renumber Vuln Scanner (now row 6)
set_cell_text(t5.rows[6].cells[0], '2.6', alignment='center')
set_cell_text(t5.rows[6].cells[3], '1', alignment='center')

# Update subtotal (now row 7)
# New cyber total: 1,700,000 + 1,200,000 + 450,000 + 240,000 + 170,000 + 180,000
cyber_total = 1_700_000 + 1_200_000 + 450_000 + 240_000 + 170_000 + 180_000
assert cyber_total == 3_940_000, f"Cyber total mismatch: {cyber_total}"
set_cell_text(t5.rows[7].cells[0],
              'รวมระบบความปลอดภัยไซเบอร์ ระยะที่ 1 (2.1–2.6)',
              bold=True, color='1A3A6B')
set_cell_text(t5.rows[7].cells[6], fmt(cyber_total), bold=True,
              color='1A3A6B', alignment='center')

print(f"  Cybersecurity phased: 6 items, subtotal {fmt(cyber_total)}")

# --- Table 6: Storage - reduce to fit ---
t6 = doc.tables[6]

# 3.1: Reduce NVMe from 500TB to 200TB, price 5.5M → 3.5M
set_cell_text(t6.rows[1].cells[1],
              'All-Flash NVMe Storage (Primary) (NetApp / Pure Storage หรือเทียบเท่า)',
              bold=True)
set_cell_text(t6.rows[1].cells[2],
              '- ความจุ Usable ≥ 200 TB\n'
              '- IOPS ≥ 1,000,000\n'
              '- NVMe-oF over InfiniBand\n'
              '- Deduplication & Compression\n'
              '- HA Dual Controller')
set_cell_text(t6.rows[1].cells[3], '1', alignment='center')
set_cell_text(t6.rows[1].cells[5], fmt(3_500_000), alignment='center')
set_cell_text(t6.rows[1].cells[6], fmt(3_500_000), bold=True,
              color='1A3A6B', alignment='center')

# 3.2: Reduce NAS from 1PB to 500TB, price 2M → 1.2M
set_cell_text(t6.rows[2].cells[1],
              'Backup & Archive Storage (NAS)',
              bold=True)
set_cell_text(t6.rows[2].cells[2],
              '- ความจุ Raw ≥ 500 TB\n'
              '- 10GbE/25GbE Network Interface\n'
              '- RAID 6 / Erasure Coding\n'
              '- Immutable Snapshot\n'
              '- Dell EMC / HPE Nimble หรือเทียบเท่า')
set_cell_text(t6.rows[2].cells[3], '1', alignment='center')
set_cell_text(t6.rows[2].cells[5], fmt(1_200_000), alignment='center')
set_cell_text(t6.rows[2].cells[6], fmt(1_200_000), bold=True,
              color='1A3A6B', alignment='center')

# 3.3: Replace Tape Library with Cloud Backup Subscription
set_cell_text(t6.rows[3].cells[0], '3.3', alignment='center')
set_cell_text(t6.rows[3].cells[1],
              'Cloud Backup & DR Subscription (3 ปี)',
              bold=True)
set_cell_text(t6.rows[3].cells[2],
              '- Cloud-based Disaster Recovery\n'
              '- Automated Backup Schedule\n'
              '- Geo-redundant Storage\n'
              '- Integration กับ On-premise Storage\n'
              '- License 3 ปี')
set_cell_text(t6.rows[3].cells[3], '1', alignment='center')
set_cell_text(t6.rows[3].cells[4], 'ชุด', alignment='center')
set_cell_text(t6.rows[3].cells[5], fmt(220_000), alignment='center')
set_cell_text(t6.rows[3].cells[6], fmt(220_000), bold=True,
              color='1A3A6B', alignment='center')

# Update storage subtotal (row 4)
storage_total = 3_500_000 + 1_200_000 + 220_000
assert storage_total == 4_920_000, f"Storage total mismatch: {storage_total}"
set_cell_text(t6.rows[4].cells[0],
              'รวมระบบ Storage & Backup (3.1–3.3)',
              bold=True, color='1A3A6B')
set_cell_text(t6.rows[4].cells[6], fmt(storage_total), bold=True,
              color='1A3A6B', alignment='center')

print(f"  Storage reduced: 3 items, subtotal {fmt(storage_total)}")

# --- Table 7: Category 1 total ---
network_total = 6_140_000
cat1_total = network_total + cyber_total + storage_total
assert cat1_total == 15_000_000, f"Cat1 total mismatch: {cat1_total}"
t7 = doc.tables[7]
set_cell_text(t7.rows[0].cells[0],
              'วงเงินรวมหมวดที่ 1  :  ระบบเครือข่ายและความปลอดภัยไซเบอร์  (รายการ 1.1–3.3)',
              bold=True, color='1A3A6B')
set_cell_text(t7.rows[0].cells[1], fmt(cat1_total), bold=True,
              color='1A3A6B', alignment='right')

print(f"  Category 1 total: {fmt(cat1_total)} ✓\n")

# --- Table 8: Category 1 notes ---
set_note_text(doc.tables[8],
    'หมายเหตุ: '
    ' | ราคา License ที่ระบุเป็นราคา 3 ปีแรก หลังจากนั้นจะจ่ายเป็นค่าต่ออายุในงบดำเนินการ '
    'ประมาณ 800,000 บาท/ปี สำหรับ SIEM, EDR และ Cloud Backup'
    ' | InfiniBand Switch และ HCA เป็นอุปกรณ์จำเป็นสำหรับ GPU Cluster '
    'High-Bandwidth Interconnect ที่มีความหน่วงต่ำ'
    ' | รายการ Cybersecurity ระยะที่ 2 (PAM, NAC, DLP, Cyber Range) '
    'วงเงินประมาณ 2,170,000 บาท จะจัดซื้อเพิ่มเติมในปีที่ 3-4 '
    'จากงบดำเนินการหรืองบสำรอง'
    ' | ใช้ Cloud Backup แทน Tape Library เพื่อลดต้นทุนฮาร์ดแวร์และบำรุงรักษา '
    'สามารถขยายความจุได้ตามความต้องการ')

# ============================================================
# STEP 3: CATEGORY 2 - Reduce H100, add Cloud GPU Credits
# Target: 65,000,000
# ============================================================
print("Step 3: Fixing Category 2 (target 65M)...")

t11 = doc.tables[11]

# 4.1: Reduce H100 from 2→1
set_cell_text(t11.rows[1].cells[3], '1', alignment='center')
set_cell_text(t11.rows[1].cells[6], fmt(14_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Add row for 4.4 Cloud GPU Credits (clone from row 3 / item 4.3)
new_tr = clone_row_after(t11, 3, 3)  # clone row 3, insert after row 3
set_tr_cell(new_tr, 0, '4.4', alignment='center')
set_tr_cell(new_tr, 1, 'Cloud GPU Credits (Burst Capacity)',
            bold=True)
set_tr_cell(new_tr, 2,
            '- NVIDIA DGX Cloud / AWS / GCP\n'
            '- สำหรับ Large-scale Experiment\n'
            '- Temporary Workload Burst\n'
            '- H100/H200 On-demand Access\n'
            '- Credit Pool 3 ปี')
set_tr_cell(new_tr, 3, '1', alignment='center')
set_tr_cell(new_tr, 4, 'ชุด', alignment='center')
set_tr_cell(new_tr, 5, fmt(4_700_000), alignment='center')
set_tr_cell(new_tr, 6, fmt(4_700_000), bold=True, color='1A3A6B',
            alignment='center')

# Re-read table after row insertion
t11 = doc.tables[11]

# Update GPU subtotal (now row 5, was row 4)
gpu_total = 14_000_000 + 18_000_000 + 7_000_000 + 4_700_000
assert gpu_total == 43_700_000, f"GPU total mismatch: {gpu_total}"
set_cell_text(t11.rows[5].cells[0],
              'รวม GPU Servers & Cloud Credits (4.1–4.4)',
              bold=True, color='1A3A6B')
set_cell_text(t11.rows[5].cells[6], fmt(gpu_total), bold=True,
              color='1A3A6B', alignment='center')

print(f"  GPU: H100 reduced to 1, Cloud Credits added. Subtotal {fmt(gpu_total)}")

# --- Table 12: Update GPU capability note ---
set_note_text(doc.tables[12],
    'รายละเอียดขีดความสามารถของ GPU Cluster รวม (4.1–4.4):'
    ' | GPU รวม: 8× H100 + 16× A100 + 16× L40S = 40 GPU (On-premise)'
    ' | รวม AI Compute Power: ≈ 4,800 TFLOPS (FP16) สำหรับ Training'
    ' | ≈ 9,600 TFLOPS (FP8) Transformer Engine'
    ' | รวม GPU VRAM: 8×80 GB (H100) + 16×80 GB (A100) + 16×48 GB (L40S) = 2,688 GB'
    ' | Cloud GPU Credits เพิ่มเติมสำหรับ Burst Capacity '
    'และ Large-scale Experiments ที่เกินกำลัง On-premise'
    ' | รองรับการ Training โมเดล LLM ขนาด 7B–70B Parameters '
    'พร้อมกันหลายโครงการ')

# Category 2 subtotals
cpu_total = 4_200_000  # unchanged
ws_sw_total = 10_200_000  # unchanged
dc_total = 6_900_000  # unchanged
cat2_total = gpu_total + cpu_total + ws_sw_total + dc_total
assert cat2_total == 65_000_000, f"Cat2 total mismatch: {cat2_total}"

# --- Table 16: Category 2 total (remove asterisk) ---
t16 = doc.tables[16]
set_cell_text(t16.rows[0].cells[0],
              'วงเงินรวมหมวดที่ 2  :  ครุภัณฑ์ AI Computing  (รายการ 4.1–7.5)',
              bold=True, color='1A3A6B')
set_cell_text(t16.rows[0].cells[1], fmt(cat2_total), bold=True,
              color='1A3A6B', alignment='right')

print(f"  Category 2 total: {fmt(cat2_total)} ✓")

# --- Table 17: Category 2 notes ---
set_note_text(doc.tables[17],
    'หมายเหตุ:'
    ' | ปรับลด GPU Server H100 จาก 2 เครื่องเหลือ 1 เครื่อง '
    'และเพิ่ม Cloud GPU Credits 4.7 ล้านบาท (3 ปี) '
    'สำหรับ Burst Capacity เพื่อให้อยู่ในกรอบวงเงิน 65 ล้านบาท'
    ' | สามารถจัดซื้อ H100 เครื่องที่ 2 ได้ในระยะที่ 2 '
    'เมื่อมีความต้องการที่ชัดเจนและงบประมาณเพิ่มเติม'
    ' | ค่าต่ออายุ License Software (NVIDIA AI Enterprise, MLOps) '
    'หลังปีที่ 3 ประมาณ 1,000,000 บาท/ปี '
    'จะจ่ายจากงบดำเนินการ'
    ' | ควรขอ Academic Discount จาก NVIDIA, Intel, AMD '
    'เนื่องจากโครงการวิจัยมหาวิทยาลัยมักได้รับส่วนลด 15–30%')

print()

# ============================================================
# STEP 4: CATEGORY 3 - Replace da Vinci, add Humanoid, adjust
# Target: 50,000,000
# ============================================================
print("Step 4: Fixing Category 3 (target 50M)...")

t24 = doc.tables[24]

# 12.6: Replace da Vinci with MicroPort Toumai trainer
set_cell_text(t24.rows[6].cells[1],
              'Surgical Robot Training Simulator (MicroPort Toumai หรือเทียบเท่า)',
              bold=True)
set_cell_text(t24.rows[6].cells[2],
              '- Surgical Simulation Platform\n'
              '- Training Module Software\n'
              '- Haptic Feedback Controller\n'
              '- 3D Visualization Screen\n'
              '- สอดคล้องกับแผนความร่วมมือจีน-อาเซียน')
set_cell_text(t24.rows[6].cells[3], '1', alignment='center')
set_cell_text(t24.rows[6].cells[5], fmt(1_500_000), alignment='center')
set_cell_text(t24.rows[6].cells[6], fmt(1_500_000), bold=True,
              color='1A3A6B', alignment='center')

# 12.7: Reduce agricultural robot from 2 to 1
set_cell_text(t24.rows[7].cells[3], '1', alignment='center')
set_cell_text(t24.rows[7].cells[6], fmt(2_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Add 12.8: Humanoid Robot (clone from row 7, insert after row 7 = before subtotal)
new_tr = clone_row_after(t24, 7, 7)
set_tr_cell(new_tr, 0, '12.8', alignment='center')
set_tr_cell(new_tr, 1,
            'Humanoid Robot Platform (Unitree G1 หรือเทียบเท่า)',
            bold=True)
set_tr_cell(new_tr, 2,
            '- Full-size Humanoid Robot\n'
            '- ≥ 23 DoF, Walking/Manipulation\n'
            '- Onboard AI Computing (Jetson/equivalent)\n'
            '- SDK: Python/ROS2 Compatible\n'
            '- สำหรับงานวิจัย Embodied AI\n'
            '  และ Human-Robot Interaction')
set_tr_cell(new_tr, 3, '2', alignment='center')
set_tr_cell(new_tr, 4, 'ตัว', alignment='center')
set_tr_cell(new_tr, 5, fmt(1_250_000), alignment='center')
set_tr_cell(new_tr, 6, fmt(2_500_000), bold=True, color='1A3A6B',
            alignment='center')

# Re-read table after insertion
t24 = doc.tables[24]

# Update Lab Equipment subtotal (now row 9, was row 8)
lab_total = (1_050_000 + 800_000 + 600_000 + 300_000 + 200_000
             + 1_500_000 + 2_000_000 + 2_500_000)
assert lab_total == 8_950_000, f"Lab total mismatch: {lab_total}"
set_cell_text(t24.rows[9].cells[0],
              'รวมอุปกรณ์สนับสนุน (12.1–12.8)',
              bold=True, color='1A3A6B')
set_cell_text(t24.rows[9].cells[6], fmt(lab_total), bold=True,
              color='1A3A6B', alignment='center')

print(f"  Lab Equipment: da Vinci→Toumai, Agri 2→1, Humanoid added. "
      f"Subtotal {fmt(lab_total)}")

# Category 3 subtotals
cobot_total = 16_200_000
auto_total = 11_100_000
sensing_total = 10_350_000
sim_total = 3_400_000
cat3_total = cobot_total + auto_total + sensing_total + sim_total + lab_total
assert cat3_total == 50_000_000, f"Cat3 total mismatch: {cat3_total}"

# --- Table 25: Category 3 total (remove asterisk) ---
t25 = doc.tables[25]
set_cell_text(t25.rows[0].cells[0],
              'วงเงินรวมหมวดที่ 3  :  ครุภัณฑ์ Robotics Lab  (รายการ 8.1–12.8)',
              bold=True, color='1A3A6B')
set_cell_text(t25.rows[0].cells[1], fmt(cat3_total), bold=True,
              color='1A3A6B', alignment='right')

print(f"  Category 3 total: {fmt(cat3_total)} ✓")

# --- Table 26: Category 3 notes ---
set_note_text(doc.tables[26],
    'หมายเหตุ:'
    ' | เพิ่มรายการ 12.8 Humanoid Robot Platform (Unitree G1) 2 ตัว '
    'เพื่อสอดคล้องกับเสาหลักด้าน Humanoid Robotics ของโครงการ '
    'และแผนความร่วมมือกับ UBTECH, Unitree และ Fourier Intelligence'
    ' | เปลี่ยนรายการ 12.6 จาก da Vinci Simulation เป็น '
    'MicroPort Toumai Surgical Trainer '
    'เพื่อสอดคล้องกับแผนพันธมิตรจีนด้าน Healthcare Robotics '
    'และต้นทุนที่ต่ำกว่า'
    ' | รายการ 12.7 Agricultural Robot Kit ลดจาก 2 ชุดเป็น 1 ชุด '
    'เป็น Flagship ที่สอดคล้องกับ Research Cluster '
    'ด้าน AI for Agriculture ของ มฟล.'
    ' | ค่าต่ออายุ License (Isaac Sim, ROS2, Isaac ROS) '
    'หลังปีที่ 3 ประมาณ 900,000 บาท/ปี จะจ่ายจากงบดำเนินการ')

print()

# ============================================================
# STEP 5: Grand total verification
# ============================================================
grand_total = cat1_total + cat2_total + cat3_total
assert grand_total == 130_000_000, f"Grand total mismatch: {grand_total}"
print(f"Grand Total: {fmt(grand_total)} ✓\n")

# ============================================================
# STEP 6: Update summary tables
# ============================================================
print("Step 5: Updating summary tables...")

# --- Table 0: Cover page metadata ---
t0 = doc.tables[0]
# Row 3: CapEx total - already says 130M, verify percentage
set_cell_text(t0.rows[3].cells[1],
              '130,000,000 บาท (ร้อยละ 43.3 ของงบโครงการทั้งหมด)',
              bold=True, color='1A1A2E')

# --- Table 27: Grand summary note ---
set_note_text(doc.tables[27],
    'ตารางสรุปวงเงินงบลงทุนรวม 3 หมวดหลัก'
    ' | งบลงทุนทั้งหมดของศูนย์ MFU-ARIC (ปีที่ 1–5) '
    'รวม 130 ล้านบาท ประกอบด้วย 3 หมวดหลัก 60 รายการครุภัณฑ์ '
    'รองรับการดำเนินงานครบทั้ง 4 Research Clusters '
    'ตั้งแต่การวิจัย การศึกษา นวัตกรรม และความร่วมมือนานาชาติ')

# --- Table 28: Summary table ---
# Structure: 17 rows × 5 cols
# [หมวด, ประเภทครุภัณฑ์, รายละเอียดย่อย, จำนวนรายการ, วงเงิน]
t28 = doc.tables[28]

# Correct item counts:
# Cat 1: Network 7 + Cyber 6 + Storage 3 = 16 items
# Cat 2: GPU 4 + CPU 2 + WS/SW 4 + DC 5 = 15 items
# Cat 3: Cobot 6 + Auto 5 + Sensing 6 + Sim 4 + Lab 8 = 29 items
# Grand total: 60 items

# Row 1: Category 1 header
set_cell_text(t28.rows[1].cells[3], '16 รายการ', bold=True, color='1A1A2E')
set_cell_text(t28.rows[1].cells[4], fmt(cat1_total), bold=True, color='1A3A6B')

# Row 2: Network sub
set_cell_text(t28.rows[2].cells[3], '7', alignment='center')
set_cell_text(t28.rows[2].cells[4], fmt(network_total), alignment='center')

# Row 3: Cybersecurity sub
set_cell_text(t28.rows[3].cells[2],
              'Firewall, SIEM, EDR, SOC (Phase 1)',
              color='1A1A2E')
set_cell_text(t28.rows[3].cells[3], '6', alignment='center')
set_cell_text(t28.rows[3].cells[4], fmt(cyber_total), alignment='center')

# Row 4: Storage sub
set_cell_text(t28.rows[4].cells[2],
              'NVMe Storage, NAS, Cloud Backup',
              color='1A1A2E')
set_cell_text(t28.rows[4].cells[3], '3', alignment='center')
set_cell_text(t28.rows[4].cells[4], fmt(storage_total), alignment='center')

# Row 5: Category 2 header
set_cell_text(t28.rows[5].cells[3], '15 รายการ', bold=True, color='1A1A2E')
set_cell_text(t28.rows[5].cells[4], fmt(cat2_total), bold=True, color='1A3A6B')

# Row 6: GPU sub (update description)
set_cell_text(t28.rows[6].cells[2],
              'H100 ×8, A100 ×16, L40S ×16, Cloud Credits',
              color='1A1A2E')
set_cell_text(t28.rows[6].cells[3], '4', alignment='center')
set_cell_text(t28.rows[6].cells[4], fmt(gpu_total), alignment='center')

# Row 7: CPU sub
set_cell_text(t28.rows[7].cells[3], '2', alignment='center')
set_cell_text(t28.rows[7].cells[4], fmt(cpu_total), alignment='center')

# Row 8: WS/SW sub
set_cell_text(t28.rows[8].cells[3], '4', alignment='center')
set_cell_text(t28.rows[8].cells[4], fmt(ws_sw_total), alignment='center')

# Row 9: DC sub
set_cell_text(t28.rows[9].cells[3], '5', alignment='center')
set_cell_text(t28.rows[9].cells[4], fmt(dc_total), alignment='center')

# Row 10: Category 3 header
set_cell_text(t28.rows[10].cells[3], '29 รายการ', bold=True, color='1A1A2E')
set_cell_text(t28.rows[10].cells[4], fmt(cat3_total), bold=True, color='1A3A6B')

# Row 11: Cobots sub (update description with updated item range)
set_cell_text(t28.rows[11].cells[3], '6', alignment='center')
set_cell_text(t28.rows[11].cells[4], fmt(cobot_total), alignment='center')

# Row 12: Autonomous sub
set_cell_text(t28.rows[12].cells[3], '5', alignment='center')
set_cell_text(t28.rows[12].cells[4], fmt(auto_total), alignment='center')

# Row 13: Sensing sub
set_cell_text(t28.rows[13].cells[3], '6', alignment='center')
set_cell_text(t28.rows[13].cells[4], fmt(sensing_total), alignment='center')

# Row 14: Sim/SW sub
set_cell_text(t28.rows[14].cells[3], '4', alignment='center')
set_cell_text(t28.rows[14].cells[4], fmt(sim_total), alignment='center')

# Row 15: Lab Equipment sub (update description and total)
set_cell_text(t28.rows[15].cells[2],
              '3D Printer, Electronics, Toumai, Agri-Robot, Humanoid',
              color='1A1A2E')
set_cell_text(t28.rows[15].cells[3], '8', alignment='center')
set_cell_text(t28.rows[15].cells[4], fmt(lab_total), alignment='center')

# Row 16: Grand total
set_cell_text(t28.rows[16].cells[0],
              'รวมงบลงทุน (Capital Expenditure) ทั้งหมด  3 หมวด  60 รายการ',
              bold=True, color='1A3A6B')
set_cell_text(t28.rows[16].cells[4], fmt(grand_total), bold=True,
              color='1A3A6B')

print("  Summary table updated with correct item counts and totals.")

# --- Table 29: Procurement phases (aligned with funding proposal) ---
t29 = doc.tables[29]
# Row 0: header [หมวด, Phase1 2570-71, Phase2 2572-73, Phase3 2574, รวม]

# Row 1: Network & Cyber
set_cell_text(t29.rows[1].cells[1], fmt(10_000_000), alignment='center')
set_cell_text(t29.rows[1].cells[2], fmt(3_000_000), alignment='center')
set_cell_text(t29.rows[1].cells[3], fmt(2_000_000), alignment='center')
set_cell_text(t29.rows[1].cells[4], fmt(15_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Row 2: AI Computing
set_cell_text(t29.rows[2].cells[1], fmt(35_000_000), alignment='center')
set_cell_text(t29.rows[2].cells[2], fmt(25_000_000), alignment='center')
set_cell_text(t29.rows[2].cells[3], fmt(5_000_000), alignment='center')
set_cell_text(t29.rows[2].cells[4], fmt(65_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Row 3: Robotics
set_cell_text(t29.rows[3].cells[1], fmt(20_000_000), alignment='center')
set_cell_text(t29.rows[3].cells[2], fmt(22_000_000), alignment='center')
set_cell_text(t29.rows[3].cells[3], fmt(8_000_000), alignment='center')
set_cell_text(t29.rows[3].cells[4], fmt(50_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Row 4: Total
set_cell_text(t29.rows[4].cells[1], fmt(65_000_000), bold=True,
              color='1A3A6B', alignment='center')
set_cell_text(t29.rows[4].cells[2], fmt(50_000_000), bold=True,
              color='1A3A6B', alignment='center')
set_cell_text(t29.rows[4].cells[3], fmt(15_000_000), bold=True,
              color='1A3A6B', alignment='center')
set_cell_text(t29.rows[4].cells[4], fmt(130_000_000), bold=True,
              color='1A3A6B', alignment='center')

# Verify phase totals
phase1 = 10_000_000 + 35_000_000 + 20_000_000
phase2 = 3_000_000 + 25_000_000 + 22_000_000
phase3 = 2_000_000 + 5_000_000 + 8_000_000
assert phase1 + phase2 + phase3 == 130_000_000, "Phase totals don't sum to 130M"
assert phase1 == 65_000_000
assert phase2 == 50_000_000
assert phase3 == 15_000_000
print("  Procurement phases aligned with funding proposal.")

# --- Table 30: Procurement strategic notes ---
set_note_text(doc.tables[30],
    'ข้อเสนอแนะเชิงกลยุทธ์ด้านการจัดซื้อจัดจ้าง'
    ' | 1. ควรดำเนินการ Vendor Evaluation อย่างเป็นระบบก่อนประกาศ TOR '
    'โดยเชิญผู้ผลิตและตัวแทนจำหน่ายมานำเสนอและทำ Demo '
    'เพื่อให้ได้ Spec ที่แม่นยำและราคาที่ดีที่สุด'
    ' | 2. GPU Server (รายการ 4.1–4.3) ควรพิจารณา Leasing Model สำหรับบางส่วน '
    'เนื่องจากเทคโนโลยีเปลี่ยนเร็ว '
    'การซื้อขาดอาจทำให้เสียโอกาสในการ Upgrade '
    'ตามวงจร Technology Refresh 3–4 ปี'
    ' | 3. ควรขอ Academic Discount จากผู้ผลิตทุกราย '
    'เนื่องจากโครงการวิจัยและการศึกษาของมหาวิทยาลัยมักได้รับส่วนลดพิเศษ 15–30% '
    '(NVIDIA, Universal Robots, ABB, Unitree มีโปรแกรม Academic Pricing ชัดเจน)'
    ' | 4. สำหรับรายการ Software License ควรเจรจา Enterprise Agreement แบบหลายปี '
    'เพื่อลดต้นทุนรวม และพิจารณา Open-source Alternatives '
    '(เช่น MinIO แทน Enterprise Storage, OpenShift แทน Commercial K8s) '
    'สำหรับส่วนที่ไม่จำเป็นต้องใช้ Commercial Support'
    ' | 5. สำคัญ: ค่าต่ออายุ Software License ทั้งหมดหลังปีที่ 3 '
    'ประมาณ 2,700,000 บาท/ปี (SIEM 400K, EDR 150K, Cloud Backup 75K, '
    'NVIDIA AI Enterprise 830K, MLOps 200K, Isaac Sim 500K, ROS2 170K, '
    'Isaac ROS 270K, DCIM 130K) '
    'ต้องจัดสรรในงบดำเนินการประจำปีตั้งแต่ปีที่ 4 เป็นต้นไป')

print("  Procurement notes updated with license renewal warning.\n")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_FILE)
print(f"=== Document saved to {OUTPUT_FILE} ===")
print(f"\nSummary of changes:")
print(f"  Category 1 (Network & Cybersecurity): {fmt(cat1_total)}")
print(f"    - Network: {fmt(network_total)} (7 items, unchanged)")
print(f"    - Cybersecurity Phase 1: {fmt(cyber_total)} (6 items, 4 deferred)")
print(f"    - Storage: {fmt(storage_total)} (3 items, reduced + cloud backup)")
print(f"  Category 2 (AI Computing): {fmt(cat2_total)}")
print(f"    - GPU + Cloud: {fmt(gpu_total)} (4 items, H100 1→1, +cloud credits)")
print(f"    - CPU: {fmt(cpu_total)} (2 items)")
print(f"    - Workstation/SW: {fmt(ws_sw_total)} (4 items)")
print(f"    - DC Infra: {fmt(dc_total)} (5 items)")
print(f"  Category 3 (Robotics Lab): {fmt(cat3_total)}")
print(f"    - Cobots: {fmt(cobot_total)} (6 items)")
print(f"    - Autonomous: {fmt(auto_total)} (5 items)")
print(f"    - Sensing: {fmt(sensing_total)} (6 items)")
print(f"    - Sim/SW: {fmt(sim_total)} (4 items)")
print(f"    - Lab Equipment: {fmt(lab_total)} (8 items, +humanoid, -da Vinci)")
print(f"  GRAND TOTAL: {fmt(grand_total)} (60 items)")
