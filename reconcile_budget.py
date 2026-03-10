"""
Reconcile budget between Equipment Specifications and Funding Proposal.
Aligns year-by-year CapEx phasing in the Proposal (T7) with the
Equipment doc's phased procurement plan (T29).

Equipment doc fixes:
  - T27: Item count 60→61, remove "Research Clusters" language
  - T29: Clarify phase headers to non-overlapping periods

Proposal fixes:
  - T7 R3 (AI Computing): Y3-4=20→25, Y5=10→5
  - T7 R4 (Robotics): Y3-4=25→22, Y5=5→8
  - T7 R5 (Network): Y1-2=8→10, Y3-4=5→3
  - T7 R13 (Total): recalculate 120/125/55 → 122/125/53
  - T7 R14/R15 (50% split): recalculate
  - T8 R3 C2: "วิจัย" → "นวัตกรรม" (missed in refocus)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

def set_cell_text(cell, text):
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')

def set_note_text(cell, text):
    parts = text.split('|')
    paras = cell.paragraphs
    for i, part in enumerate(parts):
        if i < len(paras):
            p = paras[i]
            runs = p._p.findall(qn('w:r'))
            fmt = None
            if runs:
                rPr = runs[0].find(qn('w:rPr'))
                if rPr is not None:
                    fmt = deepcopy(rPr)
            for r in runs:
                p._p.remove(r)
            new_r = etree.SubElement(p._p, qn('w:r'))
            if fmt is not None:
                new_r.insert(0, fmt)
            t = etree.SubElement(new_r, qn('w:t'))
            t.text = part.strip()
            t.set(qn('xml:space'), 'preserve')
    for i in range(len(parts), len(paras)):
        cell._tc.remove(paras[i]._p)


# ═══════════════════════════════════════════════════════════════════
# PART 1: Equipment Specifications Budget
# ═══════════════════════════════════════════════════════════════════
print("Loading Equipment Budget...")
doc_eq = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')

# ─── T27: Fix item count and language ─────────────────────────────
print("  T27: Item count 60→61, remove Research Clusters")
set_note_text(doc_eq.tables[27].rows[0].cells[0],
    "ตารางสรุปวงเงินงบลงทุนรวม 3 หมวดหลัก "
    "| งบลงทุนทั้งหมดของศูนย์ MFU-ARIC (ปีที่ 1–5) รวม 130 ล้านบาท "
    "ประกอบด้วย 3 หมวดหลัก 61 รายการครุภัณฑ์ "
    "รองรับการดำเนินงานครบทั้ง 4 องค์ประกอบหลัก "
    "(Innovation, Education, Applied Research, Regional)")

# ─── T29: Clarify phase headers ──────────────────────────────────
print("  T29: Clarify phase headers (non-overlapping)")
t29 = doc_eq.tables[29]
set_cell_text(t29.rows[0].cells[1], "ปีที่ 1–2 (พ.ศ. 2570–71) บาท")
set_cell_text(t29.rows[0].cells[2], "ปีที่ 3–4 (พ.ศ. 2572–73) บาท")
set_cell_text(t29.rows[0].cells[3], "ปีที่ 5 (พ.ศ. 2574) บาท")

doc_eq.save('MFU_ARIC_Equipment_Specifications_Budget.docx')
print("  Saved Equipment Budget.\n")


# ═══════════════════════════════════════════════════════════════════
# PART 2: Government Funding Proposal
# ═══════════════════════════════════════════════════════════════════
print("Loading Funding Proposal...")
doc_fp = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# ─── T7: Align CapEx phasing with Equipment doc T29 ──────────────
print("  T7: Align CapEx phasing with Equipment doc")
t7 = doc_fp.tables[7]

# R3 (1.2 AI Computing): 35|20|10 → 35|25|5
print("    R3 AI Computing: 35|20|10 → 35|25|5")
set_cell_text(t7.rows[3].cells[2], "25")
set_cell_text(t7.rows[3].cells[3], "5")
# R3 total stays 65

# R4 (1.3 Robotics): 20|25|5 → 20|22|8
print("    R4 Robotics: 20|25|5 → 20|22|8")
set_cell_text(t7.rows[4].cells[2], "22")
set_cell_text(t7.rows[4].cells[3], "8")
# R4 total stays 50

# R5 (1.4 Network): 8|5|2 → 10|3|2
print("    R5 Network: 8|5|2 → 10|3|2")
set_cell_text(t7.rows[5].cells[1], "10")
set_cell_text(t7.rows[5].cells[2], "3")
# R5 total stays 15

# Verify new CapEx subtotals
# Y1-2 CapEx: 30+35+20+10 = 95
# Y3-4 CapEx: 5+25+22+3 = 55
# Y5 CapEx: 0+5+8+2 = 15
# Total CapEx: 95+55+15 = 165 ✓

# OpEx (unchanged):
# Y1-2: 15+4+3+2+1+2 = 27
# Y3-4: 30+12+12+6+5+5 = 70
# Y5: 20+4+5+2+4+3 = 38
# Total OpEx: 27+70+38 = 135 ✓

# R13 (Total): 120|125|55 → 122|125|53
print("    R13 Total: 120|125|55 → 122|125|53")
set_cell_text(t7.rows[13].cells[1], "122")
# Y3-4 stays 125
set_cell_text(t7.rows[13].cells[3], "53")
# Grand total stays 300

# R14 (Gov 50%): 60|62.5|27.5 → 61|62.5|26.5
print("    R14 Gov 50%: 60|62.5|27.5 → 61|62.5|26.5")
set_cell_text(t7.rows[14].cells[1], "61")
# Y3-4 stays 62.5
set_cell_text(t7.rows[14].cells[3], "26.5")
# Total stays 150

# R15 (Other 50%): 60|62.5|27.5 → 61|62.5|26.5
print("    R15 Other 50%: 60|62.5|27.5 → 61|62.5|26.5")
set_cell_text(t7.rows[15].cells[1], "61")
# Y3-4 stays 62.5
set_cell_text(t7.rows[15].cells[3], "26.5")
# Total stays 150

# ─── T8 R3: Fix missed "วิจัย" reference ─────────────────────────
print("  T8 R3: วิจัย → นวัตกรรม")
set_cell_text(doc_fp.tables[8].rows[3].cells[2],
    "งบลงทุน: ขยายอุปกรณ์ (50%) + งบบุคลากรและนวัตกรรม (50%)")

doc_fp.save('MFU_ARIC_Government_Funding_Proposal_2569.docx')
print("  Saved Funding Proposal.\n")


# ═══════════════════════════════════════════════════════════════════
# VERIFICATION
# ═══════════════════════════════════════════════════════════════════
print("="*60)
print("VERIFICATION")
print("="*60)

# Reload and verify
doc_eq2 = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')
doc_fp2 = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# Equipment doc T29
print("\nEquipment T29 (phased budget):")
for ri, row in enumerate(doc_eq2.tables[29].rows):
    cells = [c.text.strip()[:30] for c in row.cells]
    print(f"  R{ri}: {cells}")

# Proposal T7 CapEx rows
print("\nProposal T7 (budget breakdown):")
for ri in [2, 3, 4, 5, 13, 14, 15]:
    c = doc_fp2.tables[7].rows[ri].cells
    print(f"  R{ri}: {c[0].text.strip()[:45]:45} | {c[1].text.strip():>6} | {c[2].text.strip():>6} | {c[3].text.strip():>6} | {c[4].text.strip():>6}")

# Cross-check: Proposal T7 CapEx vs Equipment T29
print("\n--- CROSS-CHECK: Proposal vs Equipment phasing ---")
# T7 values
t7 = doc_fp2.tables[7]
prop_net = [int(t7.rows[5].cells[i].text.strip()) for i in [1,2,3]]
prop_ai  = [int(t7.rows[3].cells[i].text.strip()) for i in [1,2,3]]
prop_rob = [int(t7.rows[4].cells[i].text.strip()) for i in [1,2,3]]

# T29 values (in millions, need to parse)
t29 = doc_eq2.tables[29]
def parse_m(s):
    return int(s.replace(',','').replace(' ','')) // 1000000

eq_net = [parse_m(t29.rows[1].cells[i].text) for i in [1,2,3]]
eq_ai  = [parse_m(t29.rows[2].cells[i].text) for i in [1,2,3]]
eq_rob = [parse_m(t29.rows[3].cells[i].text) for i in [1,2,3]]

print(f"  Network:  Proposal={prop_net} Equipment={eq_net} {'✓ MATCH' if prop_net==eq_net else '✗ MISMATCH'}")
print(f"  AI Comp:  Proposal={prop_ai}  Equipment={eq_ai}  {'✓ MATCH' if prop_ai==eq_ai else '✗ MISMATCH'}")
print(f"  Robotics: Proposal={prop_rob} Equipment={eq_rob} {'✓ MATCH' if prop_rob==eq_rob else '✗ MISMATCH'}")

# Verify T7 arithmetic
print("\n--- T7 ARITHMETIC ---")
y12 = sum(int(t7.rows[r].cells[1].text.strip()) for r in [2,3,4,5,7,8,9,10,11,12])
y34 = sum(int(t7.rows[r].cells[2].text.strip()) for r in [2,3,4,5,7,8,9,10,11,12])
y5  = sum(int(t7.rows[r].cells[3].text.strip().replace('–','0')) for r in [2,3,4,5,7,8,9,10,11,12])
print(f"  Y1-2: {y12} (T7 R13 says {t7.rows[13].cells[1].text.strip()}) {'✓' if str(y12)==t7.rows[13].cells[1].text.strip() else '✗'}")
print(f"  Y3-4: {y34} (T7 R13 says {t7.rows[13].cells[2].text.strip()}) {'✓' if str(y34)==t7.rows[13].cells[2].text.strip() else '✗'}")
print(f"  Y5:   {y5} (T7 R13 says {t7.rows[13].cells[3].text.strip()}) {'✓' if str(y5)==t7.rows[13].cells[3].text.strip() else '✗'}")
print(f"  Total: {y12+y34+y5} (expect 300) {'✓' if y12+y34+y5==300 else '✗'}")

# Item count
print(f"\nEquipment T27: {doc_eq2.tables[27].rows[0].cells[0].text[doc_eq2.tables[27].rows[0].cells[0].text.find('61'):doc_eq2.tables[27].rows[0].cells[0].text.find('61')+15]}")

# T8 R3 check
print(f"\nProposal T8 R3 C2: {doc_fp2.tables[8].rows[3].cells[2].text.strip()}")
