"""
enlarge_budget.py — Enlarge equipment budget 130M→150M, add DGX Spark + humanoid robots
Updates both Equipment Specs and Funding Proposal documents.

Equipment changes:
  - Cat 2 (AI Computing): +2M (DGX Spark ×10) → 65M→67M
  - Cat 3 (Robotics): +18M (4 humanoid types + dev kit) → 50M→68M
  - Grand total: 130M→150M

Proposal changes:
  - Total project: 300M→320M (gov stays 150M=47%, MFU 170M=53%)
  - T7 budget table: AI 65→67, Robotics 50→68, CapEx 165→185
  - T6 phases, T0/T1 headers, T13 summary all updated
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from copy import deepcopy


def set_cell_text(cell, text):
    """Replace cell text preserving first run's formatting."""
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        para._p.remove(r)
    new_r = etree.SubElement(para._p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')


def set_tr_cell(tr, col_idx, text):
    """Set text in a table row XML element's cell by index."""
    tcs = tr.findall(qn('w:tc'))
    if col_idx >= len(tcs):
        return
    tc = tcs[col_idx]
    p = tc.find(qn('w:p'))
    if p is None:
        return
    runs = p.findall(qn('w:r'))
    fmt = None
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            fmt = deepcopy(rPr)
    for r in runs:
        p.remove(r)
    new_r = etree.SubElement(p, qn('w:r'))
    if fmt is not None:
        new_r.insert(0, fmt)
    t = etree.SubElement(new_r, qn('w:t'))
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')


def fmt(n):
    """Format number with commas."""
    return f"{n:,}"


# ╔════════════════════════════════════════════════════════════════╗
# ║  PART 1: EQUIPMENT SPECIFICATIONS DOCUMENT                    ║
# ╚════════════════════════════════════════════════════════════════╝
print("=" * 60)
print("PART 1: Equipment Specifications")
print("=" * 60)

eq = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')

# ─────────────────────────────────────────────────────────────
# 1A. T14: Add DGX Spark row (item 6.5) before subtotal
# ─────────────────────────────────────────────────────────────
print("\n[1A] T14: Add DGX Spark (item 6.5)")
t14 = eq.tables[14]
# Clone row R4 (item 6.4) and insert after it
source_tr = t14.rows[4]._tr
new_tr = deepcopy(source_tr)
source_tr.addnext(new_tr)

set_tr_cell(new_tr, 0, "6.5")
set_tr_cell(new_tr, 1, "NVIDIA DGX Spark Desktop AI Workstation (Leadtek หรือเทียบเท่า)")
set_tr_cell(new_tr, 2,
    "- NVIDIA GB10 Grace Blackwell Superchip\n"
    "- 128 GB Unified Memory (CPU+GPU shared)\n"
    "- 1 PFLOPS FP4 AI Performance\n"
    "- DGX OS (Ubuntu-based), CUDA, TensorRT, NIM\n"
    "- Desktop form factor, พร้อมใช้งาน\n"
    "- สำหรับ Education Hub และ Startup Development")
set_tr_cell(new_tr, 3, "10")
set_tr_cell(new_tr, 4, "เครื่อง")
set_tr_cell(new_tr, 5, fmt(200_000))
set_tr_cell(new_tr, 6, fmt(2_000_000))
print("     Added: 10 × 200,000 = 2,000,000")

# Update subtotal row (now R6 after insertion, was R5)
# Subtotal: 10,200,000 + 2,000,000 = 12,200,000
set_cell_text(t14.rows[6].cells[6], fmt(12_200_000))
set_cell_text(t14.rows[6].cells[0], "รวม Workstation & Software (6.1–6.5)")
print("     Subtotal: 10,200,000 → 12,200,000")


# ─────────────────────────────────────────────────────────────
# 1B. T16: Update Category 2 total (65M → 67M)
# ─────────────────────────────────────────────────────────────
print("\n[1B] T16: Category 2 total")
set_cell_text(eq.tables[16].rows[0].cells[1], fmt(67_000_000))
print("     65,000,000 → 67,000,000")


# ─────────────────────────────────────────────────────────────
# 1C. T17: Update Category 2 notes
# ─────────────────────────────────────────────────────────────
print("\n[1C] T17: Category 2 notes")
set_cell_text(eq.tables[17].rows[0].cells[0],
    "หมายเหตุ: | ปรับลด GPU Server H100 จาก 2 เครื่องเหลือ 1 เครื่อง "
    "และเพิ่ม Cloud GPU Credits 4.7 ล้านบาท (3 ปี) สำหรับ Burst Capacity "
    "เพื่อให้อยู่ในกรอบวงเงินเดิม "
    "| เพิ่มรายการ 6.5 NVIDIA DGX Spark Desktop AI Workstation 10 เครื่อง "
    "สำหรับ Education Hub และ Startup Development "
    "(Grace Blackwell Superchip, 128 GB, 1 PFLOPS) "
    "| สามารถจัดซื้อ H100 เครื่องที่ 2 ได้ในปีที่ 3 จากงบขยายผล "
    "| วงเงินรวมหมวดที่ 2 ปรับจาก 65 ล้านบาท เป็น 67 ล้านบาท")
print("     Added DGX Spark note")


# ─────────────────────────────────────────────────────────────
# 1D. T24: Add 5 humanoid robot rows (12.9–12.13) before subtotal
# ─────────────────────────────────────────────────────────────
print("\n[1D] T24: Add humanoid robots (items 12.9–12.13)")
t24 = eq.tables[24]

# Clone row R8 (item 12.8, Unitree G1) as template
source_tr_h = t24.rows[8]._tr

# 12.9: Unitree H1
new_tr1 = deepcopy(source_tr_h)
source_tr_h.addnext(new_tr1)
set_tr_cell(new_tr1, 0, "12.9")
set_tr_cell(new_tr1, 1, "Humanoid Robot — Research Grade (Unitree H1 หรือเทียบเท่า)")
set_tr_cell(new_tr1, 2,
    "- Full-size Bipedal Humanoid, สูง ~180 cm, ~47 kg\n"
    "- ≥ 19 DoF, Walking speed ≤ 3.3 m/s\n"
    "- 3D LiDAR, Depth Camera, IMU\n"
    "- ROS2 Compatible, Python/C++ SDK\n"
    "- Battery ≥ 1 hr continuous operation\n"
    "- สำหรับวิจัย Locomotion, Manipulation, HRI")
set_tr_cell(new_tr1, 3, "2")
set_tr_cell(new_tr1, 4, "ตัว")
set_tr_cell(new_tr1, 5, fmt(3_500_000))
set_tr_cell(new_tr1, 6, fmt(7_000_000))
print("     12.9 Unitree H1: 2 × 3,500,000 = 7,000,000")

# 12.10: Fourier GR-2
new_tr2 = deepcopy(source_tr_h)
new_tr1.addnext(new_tr2)
set_tr_cell(new_tr2, 0, "12.10")
set_tr_cell(new_tr2, 1, "Humanoid Robot — Healthcare (Fourier Intelligence GR-2 หรือเทียบเท่า)")
set_tr_cell(new_tr2, 2,
    "- Healthcare Humanoid, สูง ~175 cm, ~55 kg\n"
    "- ≥ 53 DoF, Force-Torque Sensors ทุกข้อต่อ\n"
    "- Compliant Actuators สำหรับ Human Interaction\n"
    "- ROS2 Compatible, Motion Planning SDK\n"
    "- สำหรับวิจัย Healthcare, Rehabilitation, Elderly Care")
set_tr_cell(new_tr2, 3, "1")
set_tr_cell(new_tr2, 4, "ตัว")
set_tr_cell(new_tr2, 5, fmt(3_500_000))
set_tr_cell(new_tr2, 6, fmt(3_500_000))
print("     12.10 Fourier GR-2: 1 × 3,500,000 = 3,500,000")

# 12.11: UBTECH Walker S
new_tr3 = deepcopy(source_tr_h)
new_tr2.addnext(new_tr3)
set_tr_cell(new_tr3, 0, "12.11")
set_tr_cell(new_tr3, 1, "Humanoid Robot — Service (UBTECH Walker S หรือเทียบเท่า)")
set_tr_cell(new_tr3, 2,
    "- Service Humanoid, สูง ~150 cm, ~77 kg\n"
    "- ≥ 41 DoF, Dual-arm Manipulation\n"
    "- Visual SLAM, Face Recognition, NLP\n"
    "- Cloud AI Integration, Multi-language\n"
    "- สำหรับ Service Robotics, Hospitality, Customer Service")
set_tr_cell(new_tr3, 3, "1")
set_tr_cell(new_tr3, 4, "ตัว")
set_tr_cell(new_tr3, 5, fmt(2_500_000))
set_tr_cell(new_tr3, 6, fmt(2_500_000))
print("     12.11 UBTECH Walker S: 1 × 2,500,000 = 2,500,000")

# 12.12: Kepler Forerunner K2
new_tr4 = deepcopy(source_tr_h)
new_tr3.addnext(new_tr4)
set_tr_cell(new_tr4, 0, "12.12")
set_tr_cell(new_tr4, 1, "Humanoid Robot — Industrial (Kepler Forerunner K2 หรือเทียบเท่า)")
set_tr_cell(new_tr4, 2,
    "- Industrial Humanoid, สูง ~178 cm, ~85 kg\n"
    "- ≥ 40 DoF, Payload ≥ 15 kg per arm\n"
    "- High-Torque Actuators, Dexterous Hands\n"
    "- ROS2 Compatible, Industrial Communication\n"
    "- สำหรับ Industrial Automation, Manufacturing, Logistics")
set_tr_cell(new_tr4, 3, "1")
set_tr_cell(new_tr4, 4, "ตัว")
set_tr_cell(new_tr4, 5, fmt(3_000_000))
set_tr_cell(new_tr4, 6, fmt(3_000_000))
print("     12.12 Kepler K2: 1 × 3,000,000 = 3,000,000")

# 12.13: Dev Kit & Accessories
new_tr5 = deepcopy(source_tr_h)
new_tr4.addnext(new_tr5)
set_tr_cell(new_tr5, 0, "12.13")
set_tr_cell(new_tr5, 1, "Humanoid Robot Development Kit & Accessories")
set_tr_cell(new_tr5, 2,
    "- ชุดอะไหล่และอุปกรณ์สำรอง (Battery, Actuator, Sensor)\n"
    "- Charging Station & Battery Management System\n"
    "- Motion Capture Marker Set สำหรับ Humanoid\n"
    "- Maintenance Tools & Calibration Equipment\n"
    "- Training & Documentation Package")
set_tr_cell(new_tr5, 3, "1")
set_tr_cell(new_tr5, 4, "ชุด")
set_tr_cell(new_tr5, 5, fmt(2_000_000))
set_tr_cell(new_tr5, 6, fmt(2_000_000))
print("     12.13 Dev Kit: 1 × 2,000,000 = 2,000,000")

# Update subtotal row (now shifted by 5 rows)
# Find subtotal row - it's the last row in the table
last_row = t24.rows[-1]
set_cell_text(last_row.cells[0], "รวมอุปกรณ์สนับสนุน (12.1–12.13)")
# 8,950,000 + 18,000,000 = 26,950,000
set_cell_text(last_row.cells[6], fmt(26_950_000))
print("     Subtotal: 8,950,000 → 26,950,000")


# ─────────────────────────────────────────────────────────────
# 1E. T25: Update Category 3 total (50M → 68M)
# ─────────────────────────────────────────────────────────────
print("\n[1E] T25: Category 3 total")
set_cell_text(eq.tables[25].rows[0].cells[0],
    "วงเงินรวมหมวดที่ 3  :  ครุภัณฑ์ Robotics Lab  (รายการ 8.1–12.13)")
set_cell_text(eq.tables[25].rows[0].cells[1], fmt(68_000_000))
print("     50,000,000 → 68,000,000, range 8.1–12.13")


# ─────────────────────────────────────────────────────────────
# 1F. T26: Update Category 3 notes
# ─────────────────────────────────────────────────────────────
print("\n[1F] T26: Category 3 notes")
set_cell_text(eq.tables[26].rows[0].cells[0],
    "หมายเหตุ: "
    "| เพิ่มรายการ 12.8 Humanoid Robot Platform (Unitree G1) 2 ตัว เพื่อสอดคล้องกับ "
    "เสาหลักด้าน Humanoid Robotics ของโครงการ "
    "| เพิ่มรายการ 12.9–12.13 Humanoid Robot หลากหลายประเภท รวม 5 รายการ 18 ล้านบาท "
    "ครอบคลุม 4 ภาคอุตสาหกรรมเป้าหมาย: "
    "(1) Research — Unitree H1 วิจัย Locomotion/HRI, "
    "(2) Healthcare — Fourier GR-2 สำหรับ Rehabilitation, "
    "(3) Service — UBTECH Walker S สำหรับ Hospitality, "
    "(4) Industrial — Kepler K2 สำหรับ Manufacturing "
    "| สอดคล้องกับแผนความร่วมมือกับ UBTECH, Unitree, Fourier Intelligence และ Kepler Robotics "
    "| วงเงินรวมหมวดที่ 3 ปรับจาก 50 ล้านบาท เป็น 68 ล้านบาท")
print("     Added humanoid robot notes")


# ─────────────────────────────────────────────────────────────
# 1G. T27: Update summary note (130M→150M, 61→67 items)
# ─────────────────────────────────────────────────────────────
print("\n[1G] T27: Summary note")
set_cell_text(eq.tables[27].rows[0].cells[0],
    "ตารางสรุปวงเงินงบลงทุนรวม 3 หมวดหลัก\n"
    "งบลงทุนทั้งหมดของศูนย์ MFU-ARIC (ปีที่ 1–5) รวม 150 ล้านบาท "
    "ประกอบด้วย 3 หมวดหลัก 67 รายการครุภัณฑ์ "
    "รองรับการดำเนินงานครบทั้ง 4 องค์ประกอบหลัก "
    "(Innovation, Education, Applied Research, Regional)")
print("     130M→150M, 61→67 items")


# ─────────────────────────────────────────────────────────────
# 1H. T28: Update summary table
# ─────────────────────────────────────────────────────────────
print("\n[1H] T28: Summary table")
t28 = eq.tables[28]

# R5: Cat 2 header — 15→16 items, 65M→67M
set_cell_text(t28.rows[5].cells[3], "16 รายการ")
set_cell_text(t28.rows[5].cells[4], fmt(67_000_000))
print("     R5 Cat 2: 15→16 items, 65M→67M")

# R8: Workstation — (6.1-6.4)→(6.1-6.5), 4→5 items, 10.2M→12.2M
set_cell_text(t28.rows[8].cells[1], "• Workstation & Software (6.1–6.5)")
set_cell_text(t28.rows[8].cells[2], "Research WS, AI WS, DGX Spark, HPC Software")
set_cell_text(t28.rows[8].cells[3], "5")
set_cell_text(t28.rows[8].cells[4], fmt(12_200_000))
print("     R8 Workstation: +DGX Spark, 4→5, 10.2M→12.2M")

# R10: Cat 3 header — 29→34 items, 50M→68M
set_cell_text(t28.rows[10].cells[3], "34 รายการ")
set_cell_text(t28.rows[10].cells[4], fmt(68_000_000))
print("     R10 Cat 3: 29→34 items, 50M→68M")

# R15: Lab Equipment — (12.1-12.8)→(12.1-12.13), 8→13, 8.95M→26.95M
set_cell_text(t28.rows[15].cells[1], "• Lab Equipment, Humanoid Robots & Support (12.1–12.13)")
set_cell_text(t28.rows[15].cells[2],
    "3D Printer, Electronics, Toumai, Agri-Robot, "
    "Unitree G1/H1, Fourier GR-2, UBTECH Walker S, Kepler K2")
set_cell_text(t28.rows[15].cells[3], "13")
set_cell_text(t28.rows[15].cells[4], fmt(26_950_000))
print("     R15 Lab+Humanoid: 8→13, 8.95M→26.95M")

# R16: Grand total — 61→67 items, 130M→150M
set_cell_text(t28.rows[16].cells[3], "67 รายการ")
set_cell_text(t28.rows[16].cells[4], fmt(150_000_000))
print("     R16 Total: 61→67 items, 130M→150M")


# ─────────────────────────────────────────────────────────────
# 1I. T29: Update phasing table
# New: Network 10/3/2=15, AI 37/25/5=67, Robotics 25/35/8=68
# Total: 72/63/15=150
# ─────────────────────────────────────────────────────────────
print("\n[1I] T29: Phasing table")
t29 = eq.tables[29]

# R2: AI Computing — 35/25/5 → 37/25/5
set_cell_text(t29.rows[2].cells[1], fmt(37_000_000))
# C2 and C3 unchanged (25M, 5M)
set_cell_text(t29.rows[2].cells[4], fmt(67_000_000))
print("     AI Computing: 35/25/5→37/25/5 = 67M")

# R3: Robotics — 20/22/8 → 25/35/8
set_cell_text(t29.rows[3].cells[1], fmt(25_000_000))
set_cell_text(t29.rows[3].cells[2], fmt(35_000_000))
# C3 unchanged (8M)
set_cell_text(t29.rows[3].cells[4], fmt(68_000_000))
print("     Robotics: 20/22/8→25/35/8 = 68M")

# R4: Total — 65/50/15 → 72/63/15
set_cell_text(t29.rows[4].cells[1], fmt(72_000_000))
set_cell_text(t29.rows[4].cells[2], fmt(63_000_000))
# C3 unchanged (15M)
set_cell_text(t29.rows[4].cells[4], fmt(150_000_000))
print("     Total: 65/50/15→72/63/15 = 150M")

# Save equipment document
eq.save('MFU_ARIC_Equipment_Specifications_Budget.docx')
print("\n✓ Equipment Specs saved.")


# ╔════════════════════════════════════════════════════════════════╗
# ║  PART 2: FUNDING PROPOSAL DOCUMENT                            ║
# ╚════════════════════════════════════════════════════════════════╝
print("\n" + "=" * 60)
print("PART 2: Funding Proposal")
print("=" * 60)

doc = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# ─────────────────────────────────────────────────────────────
# 2A. T0: Update total amount (300M→320M)
# ─────────────────────────────────────────────────────────────
print("\n[2A] T0: Project total and gov percentage")
set_cell_text(doc.tables[0].rows[6].cells[1],
    "320,000,000 บาท (สามร้อยยี่สิบล้านบาทถ้วน)")
set_cell_text(doc.tables[0].rows[7].cells[1],
    "150,000,000 บาท (หนึ่งร้อยห้าสิบล้านบาทถ้วน) หรือประมาณร้อยละ 47")
print("     R6: 300M→320M, R7: 50%→47%")


# ─────────────────────────────────────────────────────────────
# 2B. T1: Update total amount
# ─────────────────────────────────────────────────────────────
print("\n[2B] T1 R9: Project total")
set_cell_text(doc.tables[1].rows[9].cells[1], "320,000,000 บาท")
print("     300M→320M")


# ─────────────────────────────────────────────────────────────
# 2C. T7: Update AI Computing, Robotics, totals, splits
# New totals by period:
#   Y1-2: CapEx 30+37+25+10=102, OpEx 27, Total=129
#   Y3-4: CapEx 5+25+35+3=68, OpEx 70, Total=138
#   Y5:   CapEx 0+5+8+2=15, OpEx 38, Total=53
#   Grand: 320
# ─────────────────────────────────────────────────────────────
print("\n[2C] T7: Budget table")
t7 = doc.tables[7]

# R3: AI Computing — 35/25/5=65 → 37/25/5=67
set_cell_text(t7.rows[3].cells[1], "37")
# C2, C3 unchanged
set_cell_text(t7.rows[3].cells[4], "67")
print("     R3 AI Computing: 35→37, total 65→67")

# R4: Robotics — 20/22/8=50 → 25/35/8=68
set_cell_text(t7.rows[4].cells[1], "25")
set_cell_text(t7.rows[4].cells[2], "35")
# C3 unchanged (8)
set_cell_text(t7.rows[4].cells[4], "68")
print("     R4 Robotics: 20→25/22→35, total 50→68")

# R13: Total — 122/125/53=300 → 129/138/53=320
set_cell_text(t7.rows[13].cells[1], "129")
set_cell_text(t7.rows[13].cells[2], "138")
# C3 unchanged (53)
set_cell_text(t7.rows[13].cells[4], "320")
print("     R13 Total: 122→129/125→138, grand 300→320")

# R14: Gov — 60/65/25=150 (unchanged, matches T8)
# Already correct from previous fix
print("     R14 Gov: 60/65/25=150 (unchanged)")

# R15: Other — 62/60/28=150 → 69/73/28=170
set_cell_text(t7.rows[15].cells[1], "69")
set_cell_text(t7.rows[15].cells[2], "73")
# C3 unchanged (28)
set_cell_text(t7.rows[15].cells[4], "170")
print("     R15 Other: 62→69/60→73, total 150→170")

# R14 label: update percentage
set_cell_text(t7.rows[14].cells[0],
    "– ขอรับการสนับสนุนจากภาครัฐ (ร้อยละ 47)")
print("     R14 label: 50%→47%")

# R15 label: update percentage (preserve co-funding commitment note)
set_cell_text(t7.rows[15].cells[0],
    "– มหาวิทยาลัยและแหล่งทุนอื่น (ร้อยละ 53) "
    "*จะเสนอสภามหาวิทยาลัยอนุมัติงบสมทบ "
    "และจัดทำ MOU กับแหล่งทุนก่อนเริ่มโครงการ")
print("     R15 label: 50%→53%")


# ─────────────────────────────────────────────────────────────
# 2D. T6: Update phase budgets
# Phase 1: 122→129, Phase 2: 125→138, Phase 3: 53, Total: 320
# ─────────────────────────────────────────────────────────────
print("\n[2D] T6: Phase budgets")
t6 = doc.tables[6]
set_cell_text(t6.rows[1].cells[3], "129")
set_cell_text(t6.rows[2].cells[3], "138")
# Phase 3 unchanged (53)
set_cell_text(t6.rows[4].cells[3], "320")
print("     Phase 1: 122→129, Phase 2: 125→138, Total: 300→320")


# ─────────────────────────────────────────────────────────────
# 2E. T8: Update CapEx/OpEx split description in total row
# ─────────────────────────────────────────────────────────────
print("\n[2E] T8 R6: CapEx/OpEx split")
t8 = doc.tables[8]
set_cell_text(t8.rows[6].cells[2],
    "สัดส่วนรวมทั้งโครงการ: งบลงทุน ~58% : งบดำเนินการ ~42% (โดยประมาณ)")
print("     55/45→58/42 (project-wide)")


# ─────────────────────────────────────────────────────────────
# 2F. T13: Update CapEx/OpEx percentages and equipment amount
# ─────────────────────────────────────────────────────────────
print("\n[2F] T13 R2: Budget type split")
t13 = doc.tables[13]
set_cell_text(t13.rows[2].cells[1],
    "งบลงทุน (Capital) 58%: ค่าที่ดินและสิ่งก่อสร้าง 35 ล้านบาท "
    "+ ค่าครุภัณฑ์ 150 ล้านบาท "
    "| งบดำเนินการ (Recurrent) 42%: 135 ล้านบาท "
    "| การจัดซื้อจัดจ้างครุภัณฑ์ดำเนินการตาม "
    "พ.ร.บ.การจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 "
    "รายละเอียดตามเอกสารข้อกำหนดครุภัณฑ์ (ภาคผนวก)")
print("     CapEx 55→58%, equipment 130→150M")


# ─────────────────────────────────────────────────────────────
# Save proposal document
# ─────────────────────────────────────────────────────────────
doc.save('MFU_ARIC_Government_Funding_Proposal_2569.docx')
print("\n✓ Funding Proposal saved.")


# ╔════════════════════════════════════════════════════════════════╗
# ║  VERIFICATION                                                  ║
# ╚════════════════════════════════════════════════════════════════╝
print("\n" + "=" * 60)
print("VERIFICATION")
print("=" * 60)

# Reload both docs
eq2 = Document('MFU_ARIC_Equipment_Specifications_Budget.docx')
doc2 = Document('MFU_ARIC_Government_Funding_Proposal_2569.docx')

# Equipment totals
print("\n--- Equipment Budget ---")
print(f"  Cat 1 (Network):  {eq2.tables[28].rows[1].cells[4].text.strip()}")
print(f"  Cat 2 (AI):       {eq2.tables[28].rows[5].cells[4].text.strip()}")
print(f"  Cat 3 (Robotics): {eq2.tables[28].rows[10].cells[4].text.strip()}")
print(f"  Grand Total:      {eq2.tables[28].rows[16].cells[4].text.strip()}")
print(f"  Items:            {eq2.tables[28].rows[16].cells[3].text.strip()}")

# Phasing
print("\n--- Equipment Phasing (T29) ---")
for ri in range(5):
    cells = [eq2.tables[29].rows[ri].cells[ci].text.strip()[:20] for ci in range(5)]
    print(f"  R{ri}: {cells}")

# Proposal totals
print("\n--- Proposal T7 ---")
for ri in [3,4,13,14,15]:
    cells = [doc2.tables[7].rows[ri].cells[ci].text.strip()[:40] for ci in range(5)]
    print(f"  R{ri}: {cells}")

# Proposal T0
print("\n--- Proposal T0 ---")
print(f"  Total: {doc2.tables[0].rows[6].cells[1].text.strip()}")
print(f"  Gov:   {doc2.tables[0].rows[7].cells[1].text.strip()}")
print(f"  T1 R9: {doc2.tables[1].rows[9].cells[1].text.strip()}")

# Proposal T6
print("\n--- Proposal T6 Phases ---")
for ri in range(1, 5):
    c = [doc2.tables[6].rows[ri].cells[ci].text.strip()[:30] for ci in [0,1,3]]
    print(f"  R{ri}: {c}")

# Cross-check: R14+R15=R13
print("\n--- Cross-check R14+R15=R13 ---")
t7v = doc2.tables[7]
for ci in [1,2,3,4]:
    r13 = float(t7v.rows[13].cells[ci].text.strip())
    r14 = float(t7v.rows[14].cells[ci].text.strip())
    r15 = float(t7v.rows[15].cells[ci].text.strip())
    ok = '✓' if abs(r14+r15-r13) < 0.1 else '✗'
    print(f"  Col {ci}: {r14}+{r15}={r14+r15} vs {r13} {ok}")

# T8 consistency with T7 R14
t8v = doc2.tables[8]
t8_y12 = int(t8v.rows[1].cells[1].text.strip()) + int(t8v.rows[2].cells[1].text.strip())
t8_y34 = int(t8v.rows[3].cells[1].text.strip()) + int(t8v.rows[4].cells[1].text.strip())
t8_y5 = int(t8v.rows[5].cells[1].text.strip())
r14_vals = [t7v.rows[14].cells[ci].text.strip() for ci in [1,2,3]]
print(f"\n  T8 Gov sums: Y1-2={t8_y12} Y3-4={t8_y34} Y5={t8_y5}")
print(f"  T7 R14:      Y1-2={r14_vals[0]} Y3-4={r14_vals[1]} Y5={r14_vals[2]}")
print(f"  Match: Y1-2={'✓' if str(t8_y12)==r14_vals[0] else '✗'} "
      f"Y3-4={'✓' if str(t8_y34)==r14_vals[1] else '✗'} "
      f"Y5={'✓' if str(t8_y5)==r14_vals[2] else '✗'}")

# Equipment vs Proposal consistency
eq_total = int(eq2.tables[29].rows[4].cells[4].text.strip().replace(',',''))
t7_capex_equip = (int(t7v.rows[3].cells[4].text.strip()) +
                  int(t7v.rows[4].cells[4].text.strip()) +
                  int(t7v.rows[5].cells[4].text.strip()))
print(f"\n  Equipment total: {eq_total/1e6:.0f}M")
print(f"  Proposal CapEx equipment (AI+Robot+Net): {t7_capex_equip}M")
print(f"  Match: {'✓' if eq_total/1e6 == t7_capex_equip else '✗'}")

print("\n" + "=" * 60)
print("ALL DONE")
print("=" * 60)
